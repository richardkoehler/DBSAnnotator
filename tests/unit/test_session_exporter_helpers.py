"""Broad SessionExporter helper and branch coverage (mocked I/O)."""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import pandas as pd
import pytest
from docx import Document

from dbs_annotator.models.session_data import SessionData
from dbs_annotator.utils.session_exporter import SessionExporter


@pytest.fixture
def exporter():
    sd = MagicMock(spec=SessionData)
    sd.file_path = ""
    sd.is_file_open = MagicMock(return_value=False)
    return SessionExporter(sd), sd


def test_normalize_block_id_column_variants(exporter):
    ex, _ = exporter
    df = pd.DataFrame({"block_ID": [1, 2]})
    out = ex._normalize_block_id_column(df)
    assert "block_id" in out.columns

    df2 = pd.DataFrame({"blockId": [1]})
    assert "block_id" in ex._normalize_block_id_column(df2).columns

    empty = pd.DataFrame()
    assert ex._normalize_block_id_column(empty).empty


def test_get_manufacturer_for_model(exporter):
    ex, _ = exporter
    assert ex._get_manufacturer_for_model("") == ""
    # first model in MANUFACTURERS dict
    from dbs_annotator.config_electrode_models import ELECTRODE_MODELS

    name = next(iter(ELECTRODE_MODELS))
    m = ex._get_manufacturer_for_model(name)
    assert isinstance(m, str)


def test_pick_latest_row(exporter):
    ex, _ = exporter
    assert ex._pick_latest_row(pd.DataFrame()) is None
    df = pd.DataFrame({"block_id": [1, 3, 2], "x": [1, 2, 3]})
    row = ex._pick_latest_row(df)
    assert int(row["block_id"]) == 3


def test_pick_latest_session_row(exporter):
    ex, _ = exporter
    assert ex._pick_latest_session_row(pd.DataFrame()) is None
    df = pd.DataFrame(
        {
            "session_ID": [1, 2, 2],
            "block_id": [1, 1, 2],
            "scale_value": [1, 2, 3],
        }
    )
    r = ex._pick_latest_session_row(df)
    assert r is not None


def test_column_header(exporter):
    ex, _ = exporter
    assert isinstance(ex._column_header("scale_name"), str)


def test_read_session_data_none_when_no_path(exporter):
    ex, sd = exporter
    sd.file_path = None
    assert ex._read_session_data() is None


def test_read_session_data_reads_tsv(tmp_path, exporter):
    ex, sd = exporter
    p = tmp_path / "d.tsv"
    p.write_text("a\tb\n1\t2\n", encoding="utf-8")
    sd.file_path = str(p)
    df = ex._read_session_data()
    assert df is not None
    assert len(df) == 1


def test_add_summary_section_with_notes(exporter):
    ex, _ = exporter
    doc = Document()
    df = pd.DataFrame()
    df_init = pd.DataFrame(
        {
            "session_ID": [1],
            "scale_name": ["Y"],
            "scale_value": ["1"],
            "notes": ["hello"],
            "block_id": [0],
        }
    )
    ex._add_summary_section(doc, df, df_init, df)
    assert len(doc.paragraphs) >= 1


def test_add_programming_summary_empty_df(exporter):
    ex, _ = exporter
    doc = Document()
    ex._add_programming_summary(doc, pd.DataFrame(), pd.DataFrame(), pd.DataFrame())
    assert any("No session" in p.text for p in doc.paragraphs)


def test_find_best_and_second_best_blocks_empty(exporter):
    ex, _ = exporter
    assert ex._find_best_and_second_best_blocks(pd.DataFrame()) == ([], [])


def test_find_best_and_second_best_blocks_minimal(exporter):
    ex, _ = exporter
    ex.set_scale_optimization_prefs([("Mood", "0", "10", "max", "")])
    df = pd.DataFrame(
        {
            "block_id": [1, 1],
            "scale_name": ["Mood", "Mood"],
            "scale_value": ["5", "8"],
            "laterality": ["L", "L"],
        }
    )
    a, b = ex._find_best_and_second_best_blocks(df)
    assert isinstance(a, list)
    assert isinstance(b, list)


def test_export_annotations_to_word_cancel_dialog(tmp_path, exporter):
    ex, sd = exporter
    sd.is_file_open.return_value = True
    sd.file_path = str(tmp_path / "a.tsv")
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert ex.export_annotations_to_word() is False


def test_export_annotations_to_pdf_cancel_dialog(tmp_path, exporter):
    ex, sd = exporter
    sd.is_file_open.return_value = True
    sd.file_path = str(tmp_path / "a.tsv")
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert ex.export_annotations_to_pdf() is False


def test_export_longitudinal_report_controller_calls_exporter(monkeypatch):
    from dbs_annotator.controllers.wizard_controller import WizardController

    c = WizardController()
    called = {}

    class FakeExp:
        def set_scale_optimization_prefs(self, p):
            called["prefs"] = p

        def set_clinical_scale_prefs(self, p):
            called["clinical"] = p

        def export_to_word(self, *a, **k):
            called["word"] = True

        def export_to_pdf(self, *a, **k):
            called["pdf"] = True

    def fake_exporter():
        return FakeExp()

    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.LongitudinalExporter",
        fake_exporter,
    )
    c.export_longitudinal_report([], [], "word")
    assert called.get("word")
    c.export_longitudinal_report([], [], "pdf")
    assert called.get("pdf")
