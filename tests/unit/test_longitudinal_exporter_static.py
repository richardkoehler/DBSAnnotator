"""Static helpers and API surface on LongitudinalExporter."""

from __future__ import annotations

import re
from unittest.mock import patch

from docx import Document
from PySide6.QtWidgets import QMessageBox

from dbs_annotator.utils.longitudinal_exporter import LongitudinalExporter


def test_extract_patient_id():
    assert LongitudinalExporter._extract_patient_id([]) == ""
    assert (
        LongitudinalExporter._extract_patient_id([r"C:\data\sub-ABC_task-x_events.tsv"])
        == "ABC"
    )


def test_generate_filename_with_and_without_patient():
    n = LongitudinalExporter._generate_filename(
        [r"C:\data\sub-01_task-x_events.tsv"],
        ".docx",
    )
    assert n.endswith(".docx")
    assert "sub-01" in n
    n2 = LongitudinalExporter._generate_filename([], ".pdf")
    assert n2.endswith(".pdf")
    assert re.search(r"\d{8}", n2)


def test_set_scale_prefs():
    e = LongitudinalExporter()
    e.set_scale_optimization_prefs([("Mood", "0", "10", "max", "")])
    e.set_clinical_scale_prefs([("Y", "0", "100", "min", "")])
    assert len(e.scale_optimization_prefs) == 1
    assert len(e.clinical_scale_prefs) == 1


def test_set_cell_border_top_on_table_cell():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    LongitudinalExporter._set_cell_border_top(cell)


def test_highlight_cells_empty():
    e = LongitudinalExporter()
    e._highlight_cells([], "best")


def test_add_table_legend_empty_ids():
    e = LongitudinalExporter()
    doc = Document()
    e._add_table_legend(doc, [], [])


def test_export_to_word_user_cancel(tmp_path):
    e = LongitudinalExporter()
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert e.export_to_word([str(tmp_path / "sub-01.tsv")]) is False


def test_export_to_pdf_user_cancel(tmp_path):
    e = LongitudinalExporter()
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert e.export_to_pdf([str(tmp_path / "sub-01.tsv")]) is False


def test_export_to_word_build_report_false_warns(tmp_path, monkeypatch):
    tsv = tmp_path / "sub-01_task-x_events.tsv"
    tsv.write_text("a\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    e = LongitudinalExporter()
    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.QFileDialog.getSaveFileName",
        lambda *a, **k: (str(out), "Word"),
    )
    monkeypatch.setattr(e, "_build_report", lambda *a, **k: False)
    with patch.object(QMessageBox, "warning") as w:
        assert e.export_to_word([str(tsv)], parent=None) is False
        w.assert_called()


def test_export_to_word_success(tmp_path, monkeypatch):
    tsv = tmp_path / "sub-01_task-x_events.tsv"
    tsv.write_text("a\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    e = LongitudinalExporter()
    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.QFileDialog.getSaveFileName",
        lambda *a, **k: (str(out), "Word"),
    )
    monkeypatch.setattr(e, "_build_report", lambda *a, **k: True)
    monkeypatch.setattr(e, "_show_transient_message", lambda *a, **k: None)
    assert e.export_to_word([str(tsv)], parent=None) is True


def test_export_to_pdf_success(tmp_path, monkeypatch):
    tsv = tmp_path / "sub-01_task-x_events.tsv"
    tsv.write_text("a\n", encoding="utf-8")
    pdf = tmp_path / "out.pdf"
    e = LongitudinalExporter()
    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.QFileDialog.getSaveFileName",
        lambda *a, **k: (str(pdf), "PDF"),
    )
    monkeypatch.setattr(e, "_build_report", lambda *a, **k: True)
    monkeypatch.setattr(e, "_convert_docx_to_pdf", lambda *a, **k: None)
    monkeypatch.setattr(
        LongitudinalExporter, "_open_file", staticmethod(lambda p: None)
    )
    monkeypatch.setattr(e, "_show_transient_message", lambda *a, **k: None)
    assert e.export_to_pdf([str(tsv)], parent=None) is True


def test_export_to_pdf_conversion_raises_critical(tmp_path, monkeypatch):
    tsv = tmp_path / "sub-01_task-x_events.tsv"
    tsv.write_text("a\n", encoding="utf-8")
    pdf = tmp_path / "out.pdf"
    e = LongitudinalExporter()
    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.QFileDialog.getSaveFileName",
        lambda *a, **k: (str(pdf), "PDF"),
    )
    monkeypatch.setattr(e, "_build_report", lambda *a, **k: True)

    def boom(*a, **k):
        raise RuntimeError("conv")

    monkeypatch.setattr(e, "_convert_docx_to_pdf", boom)
    with patch.object(QMessageBox, "critical") as cr:
        assert e.export_to_pdf([str(tsv)], parent=None) is False
        cr.assert_called()
