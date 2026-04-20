"""
Shared chart utilities for report generation.

Provides reusable functions for building matplotlib-based scale timeline
charts used by both the session and longitudinal exporters.  Centralising
the chart logic avoids code duplication and guarantees consistent styling
across report types.
"""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

import matplotlib
import matplotlib.pyplot as plt
import numpy as np
from docx.document import Document as DocumentType
from docx.shared import Inches
from matplotlib import cm

matplotlib.use("Agg")

logger = logging.getLogger(__name__)

# ── Colour constants (normalised 0-1 RGBA for matplotlib) ─────────
BEST_GREEN = (150 / 255, 210 / 255, 160 / 255, 160 / 255)
SECOND_GREEN = (200 / 255, 235 / 255, 205 / 255, 130 / 255)


# ── Scale-target helpers ────────────────────────────────────────────


def parse_scale_targets(
    prefs: list[tuple[str, str, str, str, str]] | None,
) -> dict[str, dict[str, Any]]:
    """Convert user preference tuples into a look-up dict.

    Args:
        prefs: [(name, min, max, mode, custom_value), ...]

    Returns:
        {scale_name: {"type": mode, "value": numeric_target}}
    """
    targets: dict[str, dict[str, Any]] = {}
    for pref in prefs or []:
        if len(pref) < 5:
            continue
        name, smin, smax, mode, custom_val = pref
        if mode == "min":
            targets[name] = {"type": "min", "value": float(smin) if smin else 0.0}
        elif mode == "max":
            targets[name] = {"type": "max", "value": float(smax) if smax else 0.0}
        elif mode == "custom":
            try:
                custom_num = float(custom_val)
            except ValueError, TypeError:
                custom_num = 0.0
            targets[name] = {"type": "custom", "value": custom_num}
    return targets


def compute_aggregate_index(
    scale_data: dict[str, dict[int, float]],
    all_points: list[int],
    scale_targets: dict[str, dict[str, Any]],
) -> dict[int, float]:
    """Compute a weighted aggregate-index value per x-point.

    The index represents how close the scales are to their respective
    targets (1.0 = perfect, 0.0 = worst observed).

    Args:
        scale_data:    {scale_name: {x_point: value}}
        all_points:    sorted list of x-axis keys
        scale_targets: output of :func:`parse_scale_targets`

    Returns:
        {x_point: index_value}
    """
    index_vals: dict[int, float] = {}
    for pt in all_points:
        weighted_scores: list[float] = []
        weights: list[float] = []

        for scale_name, pts in scale_data.items():
            if pt not in pts:
                continue
            original_value = pts[pt]

            if scale_name in scale_targets:
                info = scale_targets[scale_name]
                ttype = info["type"]
                tvalue = info["value"]

                if ttype == "min":
                    distance = original_value
                    max_possible = max(pts.values())
                    normalized = distance / max_possible if max_possible > 0 else 0
                elif ttype == "max":
                    distance = max(pts.values()) - original_value
                    max_possible = max(pts.values()) - min(pts.values())
                    normalized = distance / max_possible if max_possible > 0 else 0
                elif ttype == "custom":
                    distance = abs(original_value - tvalue)
                    max_distance = max(abs(v - tvalue) for v in pts.values())
                    normalized = distance / max_distance if max_distance > 0 else 0
                else:
                    normalized = 0.5

                weighted_scores.append(1.0 - normalized)
                weights.append(1.0)
            else:
                weighted_scores.append(0.5)
                weights.append(0.5)

        if weighted_scores and weights:
            total_w = sum(weights)
            if total_w > 0:
                index_vals[pt] = (
                    sum(w * s for w, s in zip(weights, weighted_scores, strict=False))
                    / total_w
                )
            else:
                index_vals[pt] = 0.5

    return index_vals


def find_best_and_second(
    index_vals: dict[int, float],
) -> tuple[int | None, int | None]:
    """Return the x-points with the best and second-best aggregate-index scores.

    Returns:
        (best_x, second_best_x) – either may be *None*.
    """
    if not index_vals:
        return None, None
    ranked = sorted(index_vals, key=lambda k: index_vals[k], reverse=True)
    best = ranked[0]
    second = ranked[1] if len(ranked) > 1 else None
    return best, second


# ── Chart rendering ────────────────────────────────────────────────

# Matplotlib line styles matching the old set
_LINE_STYLES = ["-", "--", ":", "-.", (0, (3, 1, 1, 1, 1, 1))]


def build_scales_chart(
    scale_data: dict[str, dict[int, float]],
    scale_prefs: list[tuple[str, str, str, str, str]] | None,
    *,
    title: str = "",
    x_label: str = "X",
    y_label: str = "Scale Value",
    x_ticks: list[tuple[int, str]] | None = None,
    width: int = 1100,
    height: int = 520,
    show_general_index: bool = True,
    rotate_x_ticks: bool = False,
) -> bytes | None:
    """Build a scale-trend chart and return it as PNG bytes.

    The chart includes:
    * Rainbow-coloured individual scale lines
    * A thick black Aggregate Index line (when *show_general_index* and ≥ 2 scales)
    * Green vertical bands for best (dark) and second-best (light)
    * A compact legend above the plot

    Args:
        scale_data:          {scale_name: {x_point: value}}
        scale_prefs:         user optimisation preferences (may be *None*)
        title:               chart title text (empty string for no title)
        x_label:             bottom-axis label
        y_label:             left-axis label
        x_ticks:             optional custom bottom-axis ticks
        width:               image width in px
        height:              image height in px
        show_general_index:  whether to draw the Aggregate Index line
        rotate_x_ticks:      whether to rotate x-axis tick labels by 90 degrees

    Returns:
        PNG image bytes, or *None* on failure.
    """
    try:
        n_scales = len(scale_data)
        if n_scales == 0:
            return None

        dpi = 100
        fig_w = width / dpi
        # Increase height when rotating ticks
        actual_height = height + 150 if rotate_x_ticks else height
        fig_h = actual_height / dpi

        cmap_obj = cm.get_cmap("Dark2")
        colors = [cmap_obj(i % cmap_obj.N)[:3] for i in range(n_scales)]

        has_index = show_general_index and n_scales >= 2

        fig, ax1 = plt.subplots(figsize=(fig_w, fig_h), dpi=dpi)
        fig.patch.set_facecolor("white")
        ax1.set_facecolor("white")

        # ── Collect all x-points for NaN gap handling ─────────────
        all_x = sorted({x for pts in scale_data.values() for x in pts})

        # ── Plot individual scales ──────────────────────────────────
        for idx, (sname, pts) in enumerate(scale_data.items()):
            c = colors[idx]
            ls = _LINE_STYLES[idx % len(_LINE_STYLES)]
            ys = [pts.get(x, float("nan")) for x in all_x]
            # Build segment arrays so NaN creates gaps in lines
            xs_arr = np.array(all_x, dtype=float)
            ys_arr = np.array(ys, dtype=float)
            ax1.plot(
                xs_arr,
                ys_arr,
                color=c,
                linewidth=2,
                linestyle=ls,
                marker="o",
                markersize=6,
                markerfacecolor=c,
                markeredgecolor=c,
                markeredgewidth=1,
                label=sname,
            )

        # ── Aggregate Index on right y-axis ────────────────────────
        best_x: int | None = None
        second_x: int | None = None
        ax2 = None
        if has_index:
            all_points = sorted({x for pts in scale_data.values() for x in pts})
            scale_targets = parse_scale_targets(scale_prefs)
            index_vals = compute_aggregate_index(scale_data, all_points, scale_targets)

            if index_vals:
                ax2 = ax1.twinx()
                ix = sorted(index_vals.keys())
                iy = [index_vals[x] for x in ix]
                ax2.plot(
                    ix,
                    iy,
                    color="black",
                    linewidth=3,
                    marker="D",
                    markersize=7,
                    markerfacecolor="black",
                    markeredgecolor="black",
                    label="Aggregate Index",
                    zorder=5,
                )
                ax2.set_ylim(0, 1)
                ax2.set_ylabel(
                    "Aggregate Index Score",
                    fontsize=12,
                    fontfamily="Arial",
                    color="black",
                )
                ax2.tick_params(axis="y", labelsize=10)
                best_x, second_x = find_best_and_second(index_vals)

        # ── Best / second-best green vertical bands ─────────────────
        if best_x is not None:
            ax1.axvspan(best_x - 0.35, best_x + 0.35, color=BEST_GREEN, zorder=0)
        if second_x is not None and second_x != best_x:
            ax1.axvspan(second_x - 0.35, second_x + 0.35, color=SECOND_GREEN, zorder=0)

        # ── Axes labels and styling ─────────────────────────────────
        if title:
            ax1.set_title(title, fontsize=16, fontfamily="Arial", color="black")
        ax1.set_ylabel(y_label, fontsize=12, fontfamily="Arial", color="black")
        ax1.set_xlabel(x_label, fontsize=12, fontfamily="Arial", color="black")
        ax1.tick_params(axis="both", labelsize=10)
        ax1.grid(True, alpha=0.3)

        # ── X-ticks ─────────────────────────────────────────────────
        if x_ticks is not None:
            tick_positions = [t[0] for t in x_ticks]
            tick_labels = [t[1] for t in x_ticks]
            ax1.set_xticks(tick_positions)
            if rotate_x_ticks:
                ax1.set_xticklabels(
                    tick_labels,
                    rotation=90,
                    ha="center",
                    fontsize=10,
                )
            else:
                ax1.set_xticklabels(tick_labels, fontsize=10)

        # Ensure x-range includes all tick positions with padding
        if x_ticks:
            x_min = min(t[0] for t in x_ticks) - 0.5
            x_max = max(t[0] for t in x_ticks) + 0.5
            ax1.set_xlim(x_min, x_max)

        # ── Legend ──────────────────────────────────────────────────
        handles1, labels1 = ax1.get_legend_handles_labels()
        handles2, labels2 = ([], [])
        if ax2 is not None:
            handles2, labels2 = ax2.get_legend_handles_labels()
        all_handles = handles1 + handles2
        all_labels = labels1 + labels2
        if all_handles:
            n_cols = max(1, len(all_handles))
            fig.legend(
                all_handles,
                all_labels,
                loc="upper center",
                ncol=n_cols,
                fontsize=9,
                frameon=True,
                facecolor="white",
                edgecolor=(0.6, 0.6, 0.6, 0.4),
                framealpha=0.9,
                bbox_to_anchor=(0.5, 1.0),
            )

        fig.tight_layout(rect=(0, 0, 1, 0.93))

        # ── Export to PNG bytes ─────────────────────────────────────
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        png_bytes = buf.read()
        buf.close()
        return png_bytes

    except Exception as exc:
        logger.exception("build_scales_chart failed: %s", exc)
        return None


def add_chart_to_doc(
    doc: DocumentType,
    png_bytes: bytes | None,
    *,
    heading: str | None = None,
    heading_level: int = 2,
    width_inches: float | None = None,
    fallback_message: str = "Chart generation error.",
) -> None:
    """Insert a PNG chart into a Word document.

    Args:
        doc:              python-docx Document instance
        png_bytes:        raw PNG bytes (or *None* on failure)
        heading:          optional heading text above the chart
        heading_level:    heading level (default 2)
        width_inches:     image width in the document (None for full page width)
        fallback_message: text shown if *png_bytes* is None
    """
    if heading:
        doc.add_heading(heading, level=heading_level)

    # Calculate page width if not specified
    if width_inches is None:
        section = doc.sections[0]
        page_w = (
            int(section.page_width or 0)
            - int(section.left_margin or 0)
            - int(section.right_margin or 0)
        ) / 914400  # Convert twips to inches
        width_inches = max(4.0, page_w)  # Minimum 4 inches, otherwise full page

    if png_bytes is None:
        doc.add_paragraph(fallback_message)
        return

    img_buf = BytesIO(png_bytes)
    doc.add_picture(img_buf, width=Inches(width_inches))
    doc.add_paragraph()
    img_buf.close()
