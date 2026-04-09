"""
Longitudinal report exporter for Clinical DBS Annotator.

Combines data from multiple annotation TSV files and generates a unified
longitudinal report in Word or PDF format, with best-entry highlighting
based on user-selected scale optimization preferences.
"""

import os
import re
import tempfile
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PySide6.QtCore import Qt, QTimer
from PySide6.QtWidgets import QFileDialog, QMessageBox, QWidget

from .. import __app_name__, __version__
from ..config import PLACEHOLDERS
from ..config_electrode_models import ELECTRODE_MODELS, MANUFACTURERS, ContactState


class LongitudinalExporter:
    """Generate longitudinal reports from multiple annotation TSV files."""

    def __init__(self):
        self.scale_optimization_prefs: list = []

    def set_scale_optimization_prefs(self, prefs: list) -> None:
        """Set scale optimization preferences for best-entry highlighting."""
        self.scale_optimization_prefs = prefs or []

    # ------------------------------------------------------------------
    # Public export API
    # ------------------------------------------------------------------

    def export_to_word(
        self, file_paths: list[str], parent: QWidget | None = None, sections=None
    ) -> bool:
        """Export longitudinal report to Word format."""
        try:
            default_name = self._generate_filename(file_paths, ".docx")
            start_dir = os.path.dirname(file_paths[0]) if file_paths else ""
            start_path = os.path.join(start_dir, default_name) if start_dir else default_name

            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Longitudinal Report",
                start_path,
                "Word Files (*.docx);;All Files (*)",
            )
            if not file_path:
                return False
            if not file_path.endswith(".docx"):
                file_path += ".docx"

            ok = self._build_report(file_paths, file_path, sections=sections)
            if not ok:
                QMessageBox.warning(
                    parent, "No Data", "No session data found in the loaded files."
                )
                return False

            self._show_transient_message(
                parent, "Export Completed", f"Report saved:\n{file_path}"
            )
            return True
        except Exception as e:
            QMessageBox.critical(
                parent, "Export Error", f"Failed to export report:\n{e}"
            )
            return False

    def export_to_pdf(
        self, file_paths: list[str], parent: QWidget | None = None, sections=None
    ) -> bool:
        """Export longitudinal report to PDF (via intermediate Word)."""
        try:
            default_name = self._generate_filename(file_paths, ".pdf")
            start_dir = os.path.dirname(file_paths[0]) if file_paths else ""
            start_path = os.path.join(start_dir, default_name) if start_dir else default_name

            pdf_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Longitudinal Report",
                start_path,
                "PDF Files (*.pdf);;All Files (*)",
            )
            if not pdf_path:
                return False
            if not pdf_path.endswith(".pdf"):
                pdf_path += ".pdf"

            docx_tmp = os.path.splitext(pdf_path)[0] + "_tmp.docx"
            ok = self._build_report(file_paths, docx_tmp, sections=sections)
            if not ok:
                QMessageBox.warning(
                    parent, "No Data", "No session data found in the loaded files."
                )
                return False

            try:
                self._convert_docx_to_pdf(docx_tmp, pdf_path)
            finally:
                try:
                    os.unlink(docx_tmp)
                except Exception:
                    pass

            self._show_transient_message(
                parent, "Export Completed", f"Report saved:\n{pdf_path}"
            )
            return True
        except Exception as e:
            QMessageBox.critical(
                parent, "Export Error", f"Failed to export report:\n{e}"
            )
            return False

    # ------------------------------------------------------------------
    # Report building
    # ------------------------------------------------------------------

    def _build_report(self, file_paths: list[str], out_path: str, sections=None) -> bool:
        """Read all files, merge, and build the Word document."""
        # Sort files chronologically by earliest date+time in each file
        def get_file_datetime(path):
            try:
                df = pd.read_csv(path, sep="\t")
                if "date" in df.columns and "time" in df.columns:
                    # Combine date and time to create datetime for sorting
                    df["datetime"] = pd.to_datetime(df["date"] + " " + df["time"], errors="coerce")
                    valid_times = df["datetime"].dropna()
                    if not valid_times.empty:
                        return valid_times.min()  # Use earliest time in file
                # Fallback to filename date if available
                basename = os.path.basename(path)
                import re
                date_match = re.search(r'ses-(\d{8})', basename)
                if date_match:
                    date_str = date_match.group(1)
                    return pd.to_datetime(date_str, format="%Y%m%d")
                # If no date info available, return a very old date to put it at the end
                return pd.Timestamp("1900-01-01")
            except Exception:
                return pd.Timestamp("1900-01-01")

        # Sort files from oldest to newest
        file_paths = sorted(file_paths, key=get_file_datetime)

        frames = []
        for path in file_paths:
            try:
                df = pd.read_csv(path, sep="\t")
                # Tag each row with its source file for traceability
                df["_source_file"] = os.path.basename(path)
                frames.append(df)
            except Exception as e:
                print(f"[WARNING] Could not read {path}: {e}")

        if not frames:
            return False

        df_all = pd.concat(frames, ignore_index=True)
        if df_all.empty:
            return False

        df_all = self._normalize_block_id(df_all)

        # Split initial vs session rows
        if "is_initial" in df_all.columns:
            df_all["is_initial"] = (
                pd.to_numeric(df_all["is_initial"], errors="coerce").fillna(0).astype(int)
            )
            df_all[df_all["is_initial"] == 1]
            df_session = df_all[df_all["is_initial"] == 0]
        else:
            df_all.iloc[0:0]
            df_session = df_all

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        # Title
        title = doc.add_heading("Longitudinal DBS Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Generated on: {datetime.now().astimezone().strftime('%Y-%m-%d %H:%M:%S')}"
            f" by {__app_name__} v{__version__}"
        )

        # Patient info (from first file)
        patient_id = self._extract_patient_id(file_paths)
        if patient_id:
            doc.add_paragraph(f"Patient ID: {patient_id}")

        doc.add_paragraph(f"Files included: {len(file_paths)}")
        for fp in file_paths:
            doc.add_paragraph(f"  {os.path.basename(fp)}")

        doc.add_paragraph("")

        # Determine which sections to include (default: sessions_overview + session_data)
        all_keys = ["sessions_overview", "session_data", "electrode_config", "programming_summary"]
        active = set(sections) if sections is not None else {"sessions_overview", "session_data"}

        # Render in the defined order
        for key in all_keys:
            if key not in active:
                continue
            if key == "sessions_overview":
                self._add_sessions_overview(doc, df_all, file_paths)
                self._add_scales_timeline_chart(doc, df_session)
                doc.add_paragraph("")
            elif key == "session_data":
                self._add_scales_timeline_chart(doc, df_session)
                self._add_longitudinal_data_table(doc, df_session, file_paths)
                doc.add_paragraph("")
            elif key == "electrode_config":
                self._add_electrode_config_section(doc, df_all, file_paths)
                doc.add_paragraph("")
            elif key == "programming_summary":
                self._add_programming_summary(doc, df_all, file_paths)
                doc.add_paragraph("")

        doc.save(out_path)
        return True

    # ------------------------------------------------------------------
    # Report sections
    # ------------------------------------------------------------------

    def _add_sessions_overview(
        self, doc: Document, df: pd.DataFrame, file_paths: list[str]
    ) -> None:
        """Add a summary table listing each session file with date and entry count."""
        doc.add_heading("Sessions Overview", level=1)

        headers = ["#", "File", "Date", "Entries", "Clinical scales", "Values"]
        col_widths = [0.25, 2.0, 1.0, 0.7, 1.2, 0.8]  # inches
        table = doc.add_table(rows=1 + len(file_paths), cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False

        w_twips = [Inches(w) for w in col_widths]
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = w_twips[ci]

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for idx, fp in enumerate(file_paths):
            row_cells = table.rows[idx + 1].cells
            basename = os.path.basename(fp)
            row_cells[0].text = str(idx + 1)
            row_cells[1].text = basename

            sub_df = df[df["_source_file"] == basename] if "_source_file" in df.columns else df
            date_str = ""
            if "date" in sub_df.columns and not sub_df.empty:
                dates = sub_df["date"].dropna().unique()
                if len(dates) > 0:
                    date_str = ", ".join(str(d) for d in sorted(dates))
            row_cells[2].text = date_str

            session_rows = sub_df
            if "is_initial" in sub_df.columns:
                session_rows = sub_df[sub_df["is_initial"] == 0]

            # Count unique block_ids (entries)
            if "block_id" in session_rows.columns:
                unique_entries = session_rows["block_id"].nunique()
            else:
                unique_entries = len(session_rows)
            row_cells[3].text = str(unique_entries)

            # Collect scales from is_initial=1 rows with highest block_id per file
            scale_pairs = []
            if "scale_name" in sub_df.columns and "scale_value" in sub_df.columns:
                # Filter for is_initial=1 only (baseline)
                baseline_df = sub_df.copy()
                if "is_initial" in baseline_df.columns:
                    baseline_df = baseline_df[pd.to_numeric(baseline_df["is_initial"], errors="coerce").fillna(0).astype(int) == 1]

                if not baseline_df.empty and "block_id" in baseline_df.columns:
                    try:
                        baseline_df["block_id_num"] = pd.to_numeric(baseline_df["block_id"], errors="coerce")
                        max_block = baseline_df["block_id_num"].max()

                        # Get ALL rows with the highest block_id (there could be multiple)
                        max_block_rows = baseline_df[baseline_df["block_id_num"] == max_block]

                        # Collect all scale pairs from these rows
                        all_scales = {}
                        for _, row in max_block_rows.iterrows():
                            sn = str(row.get("scale_name", "") or "").strip()
                            sv = str(row.get("scale_value", "") or "").strip()

                            if sn and sv:
                                sn_lines = [s.strip() for s in sn.split("\n") if s.strip()]
                                sv_lines = [s.strip() for s in sv.split("\n") if s.strip()]

                                # Store scales, keeping first non-NaN value per scale name
                                for name, val in zip(sn_lines, sv_lines, strict=False):
                                    if val != "NaN" and val.strip() != "NaN":
                                        if name not in all_scales:
                                            all_scales[name] = val

                        # Convert to list of tuples
                        scale_pairs = list(all_scales.items())
                    except Exception:
                        scale_pairs = []

            row_cells[4].text = "\n".join(p[0] for p in scale_pairs) if scale_pairs else ""
            row_cells[5].text = "\n".join(p[1] for p in scale_pairs) if scale_pairs else ""

    def _add_electrode_config_section(
        self, doc: Document, df_all: pd.DataFrame, file_paths: list[str]
    ) -> None:
        """Add per-file electrode configuration (Initial / Final, Left / Right).

        Each file gets its own heading and a 4-column table matching the
        single-session report layout.  A page break separates consecutive files.
        """
        if df_all is None or df_all.empty:
            return
        if "electrode_model" not in df_all.columns:
            return
        if "_source_file" not in df_all.columns:
            return

        doc.add_heading("Electrode Configurations", level=1)

        from docx.enum.text import WD_BREAK

        any_rendered = False
        for _fp_idx, fp in enumerate(file_paths):
            basename = os.path.basename(fp)
            sub = df_all[df_all["_source_file"] == basename].copy()
            if sub.empty:
                continue

            if "is_initial" in sub.columns:
                sub["is_initial"] = pd.to_numeric(sub["is_initial"], errors="coerce").fillna(0).astype(int)
                df_init = sub[sub["is_initial"] == 1]
                df_final = sub[sub["is_initial"] == 0]
            else:
                df_init = sub.iloc[0:0]
                df_final = sub

            if df_init.empty and df_final.empty:
                continue

            # Pick representative rows
            init_row = self._pick_latest_row(df_init) if not df_init.empty else None
            final_row = self._pick_latest_row(df_final) if not df_final.empty else None

            model_name = ""
            for candidate in (final_row, init_row):
                if candidate is not None:
                    m = str(candidate.get("electrode_model", "") or "").strip()
                    if m:
                        model_name = m
                        break
            if not model_name:
                continue

            # Page break before each file except the first
            if any_rendered:
                para = doc.add_paragraph()
                para.add_run().add_break(WD_BREAK.PAGE)

            any_rendered = True

            # File sub-heading
            label = basename.replace("_events.tsv", "").replace(".tsv", "")
            doc.add_heading(label, level=2)

            manufacturer = self._get_manufacturer_for_model(model_name)
            if manufacturer:
                doc.add_paragraph(f"Electrode model: {manufacturer} | {model_name}")
            else:
                doc.add_paragraph(f"Electrode model: {model_name}")

            # Helper to extract contact strings from a row
            def _contacts(row, side):
                if row is None:
                    return "", ""
                anode = str(row.get(f"{side}_anode", "") or "")
                cathode = str(row.get(f"{side}_cathode", "") or "")
                return anode, cathode

            i_la, i_lc = _contacts(init_row, "left")
            i_ra, i_rc = _contacts(init_row, "right")
            f_la, f_lc = _contacts(final_row, "left")
            f_ra, f_rc = _contacts(final_row, "right")

            # Render PNGs
            tmp_files = []
            init_model = str(init_row.get("electrode_model", "") or model_name) if init_row is not None else model_name
            final_model = str(final_row.get("electrode_model", "") or model_name) if final_row is not None else model_name

            png_init_l = self._render_electrode_png(init_model, i_la, i_lc) if init_row is not None else None
            png_init_r = self._render_electrode_png(init_model, i_ra, i_rc) if init_row is not None else None
            png_final_l = self._render_electrode_png(final_model, f_la, f_lc) if final_row is not None else None
            png_final_r = self._render_electrode_png(final_model, f_ra, f_rc) if final_row is not None else None

            for p in (png_init_l, png_init_r, png_final_l, png_final_r):
                if p:
                    tmp_files.append(p)

            # 4-column x 4-row table: Init L | Init R | Final L | Final R
            t = doc.add_table(rows=4, cols=4)
            t.autofit = False

            # Remove borders
            tbl = t._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr() # noqa: N806
            borders = OxmlElement("w:tblBorders")
            for bname in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{bname}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                borders.append(b)
            tblPr.append(borders)

            # Row 0: "Initial Settings" (merged 0-1) | "Final Settings" (merged 2-3)
            for merged_start, merged_end, heading_text in [(0, 1, "Initial Settings"), (2, 3, "Final Settings")]:
                cell = t.cell(0, merged_start).merge(t.cell(0, merged_end))
                cell.text = heading_text
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True

            # Row 1: Left | Right | Left | Right
            for col, txt in enumerate(["Left", "Right", "Left", "Right"]):
                t.cell(1, col).text = txt
                for p in t.cell(1, col).paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Row 2: config text
            configs = [
                (i_la, i_lc),
                (i_ra, i_rc),
                (f_la, f_lc),
                (f_ra, f_rc),
            ]
            for col, (anode, cathode) in enumerate(configs):
                t.cell(2, col).text = f"+{anode}\n-{cathode}".strip()
                for p in t.cell(2, col).paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Row 3: electrode images
            png_list = [png_init_l, png_init_r, png_final_l, png_final_r]
            for col, png in enumerate(png_list):
                cell = t.cell(3, col)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if png:
                    run = p.add_run()
                    try:
                        run.add_picture(png, width=Inches(1.15))
                    except Exception:
                        pass

            # Cleanup temp PNG files
            for pth in tmp_files:
                try:
                    os.unlink(pth)
                except Exception:
                    pass

            doc.add_paragraph("")

    def _add_programming_summary(
        self, doc: Document, df_all: pd.DataFrame, file_paths: list[str]
    ) -> None:
        """Add a per-session programming summary table."""
        if df_all is None or df_all.empty:
            return
        if "_source_file" not in df_all.columns:
            return

        doc.add_heading("Programming Summary", level=1)

        headers = ["Session", "Configurations", "Amplitude (L)", "Amplitude (R)",
                    "Frequency (L)", "Frequency (R)", "Pulse Width (L)", "Pulse Width (R)"]
        rows_data = []

        for fp in file_paths:
            basename = os.path.basename(fp)
            sub = df_all[df_all["_source_file"] == basename]
            if sub.empty:
                continue

            if "is_initial" in sub.columns:
                sub_sess = sub[pd.to_numeric(sub["is_initial"], errors="coerce").fillna(0).astype(int) == 0]
            else:
                sub_sess = sub

            label = basename.replace("_events.tsv", "").replace(".tsv", "")

            sub_n = self._normalize_block_id(sub_sess)
            n_configs = sub_n["block_id"].nunique() if "block_id" in sub_n.columns else 0

            def _range_str(series, unit=""):
                vals = pd.to_numeric(series, errors="coerce").dropna()
                if vals.empty:
                    return "N/A"
                mn, mx = vals.min(), vals.max()
                if mn == mx:
                    return f"{mn:.1f}{unit}" if unit else f"{mn:g}"
                return f"{mn:.1f}–{mx:.1f}{unit}" if unit else f"{mn:g}–{mx:g}"

            amp_l = _range_str(sub_sess.get("left_amplitude", pd.Series()), " mA")
            amp_r = _range_str(sub_sess.get("right_amplitude", pd.Series()), " mA")
            freq_l = _range_str(sub_sess.get("left_stim_freq", pd.Series()), " Hz")
            freq_r = _range_str(sub_sess.get("right_stim_freq", pd.Series()), " Hz")
            pw_l = _range_str(sub_sess.get("left_pulse_width", pd.Series()), " µs")
            pw_r = _range_str(sub_sess.get("right_pulse_width", pd.Series()), " µs")

            rows_data.append([label, str(n_configs), amp_l, amp_r, freq_l, freq_r, pw_l, pw_r])

        if not rows_data:
            doc.add_paragraph("No programming data available.")
            return

        table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
        table.style = "Table Grid"

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for r_idx, row_vals in enumerate(rows_data):
            for c_idx, val in enumerate(row_vals):
                table.rows[r_idx + 1].cells[c_idx].text = val

        doc.add_paragraph("")

    def _add_longitudinal_data_table(
        self, doc: Document, df_session: pd.DataFrame,
        file_paths: list[str] | None = None,
    ) -> None:
        """Add the main longitudinal data table with green highlighting."""
        doc.add_heading("Session Data", level=1)

        if df_session is None or df_session.empty:
            doc.add_paragraph("No session data available.")
            return

        lateral_df = self._create_lateral_table(df_session)
        if lateral_df.empty:
            doc.add_paragraph("No session data available.")
            return

        # Ensure date column is present (populated inside _create_lateral_table)
        if "date" not in lateral_df.columns:
            lateral_df["date"] = ""

        columns_to_exclude = [
            "time", "onset", "block_id", "session_ID", "source",
            "is_initial", "electrode_model", "_source_file", "_global_entry_id",
        ]
        display_cols = [c for c in lateral_df.columns if c not in columns_to_exclude]

        lateral_cols = ["date", "laterality", "frequency", "anode", "cathode", "amplitude", "pulse_width"]
        common_cols = ["program_ID", "scale_name", "scale_value", "notes"]

        lateral_cols = [c for c in lateral_cols if c in display_cols]
        common_cols = [c for c in common_cols if c in display_cols]
        ordered = lateral_cols + common_cols

        if not ordered:
            doc.add_paragraph("No displayable columns found.")
            return

        table = doc.add_table(rows=lateral_df.shape[0] + 1, cols=len(ordered))
        table.style = "Table Grid"
        table.autofit = False

        # Column widths
        section = doc.sections[0]
        page_w = (section.page_width - section.left_margin - section.right_margin) / 914400
        base_w = {
            "date": 0.65, "laterality": 0.25, "program_ID": 0.35,
            "frequency": 0.45, "anode": 0.45, "cathode": 0.60,
            "amplitude": 0.60, "pulse_width": 0.50,
            "scale_name": 1.00, "scale_value": 0.55,
        }
        widths = [base_w.get(c, 0.5) for c in ordered]
        if "notes" in ordered:
            ni = ordered.index("notes")
            used = sum(w for j, w in enumerate(widths) if j != ni)
            widths[ni] = max(1.5, page_w - used)

        w_twips = [Inches(max(0.25, w)) for w in widths]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = w_twips[idx]

        # Header
        for i, col in enumerate(ordered):
            cell = table.rows[0].cells[i]
            cell.text = self._column_header(col)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        # Find best / second-best entries
        best_ids, second_ids = self._find_best_and_second_best(lateral_df)

        prev_entry_id = None
        for i, (_, row) in enumerate(lateral_df.iterrows()):
            row_cells = table.rows[i + 1].cells

            entry_id = row.get("_global_entry_id", None)
            if best_ids and entry_id in best_ids:
                self._highlight_cells(row_cells, "best")
            elif second_ids and entry_id in second_ids:
                self._highlight_cells(row_cells, "second")

            # Separator between entries
            if (
                prev_entry_id is not None
                and entry_id != prev_entry_id
                and row.get("laterality") == "L"
            ):
                for cell in row_cells:
                    self._set_cell_border_top(cell, sz=24)
            prev_entry_id = entry_id

            for j, col in enumerate(ordered):
                val = row.get(col, "")
                cell_text = str(val) if pd.notna(val) else ""
                if col in ("frequency", "pulse_width"):
                    try:
                        v = float(val)
                        cell_text = str(int(v)) if v == int(v) else str(v)
                    except (ValueError, TypeError):
                        pass

                if col in common_cols and row.get("laterality") == "R" and i > 0:
                    prev_cell = table.rows[i].cells[j]
                    prev_cell.merge(row_cells[j])
                    row_cells[j].text = ""
                elif col == "cathode" and "_" in cell_text:
                    # Multi-contact cathode: show stacked with Total label
                    contacts = cell_text.replace("_", "\n")
                    row_cells[j].text = contacts + "\nTotal"
                elif col == "amplitude" and "_" in cell_text:
                    # Multi-contact amplitude: show stacked values with total
                    parts = cell_text.split("_")
                    try:
                        # Validate all parts are numbers and calculate total
                        values = [float(p) for p in parts]
                        total = sum(values)
                        total_str = f"{total:.2f}".rstrip("0").rstrip(".")
                        row_cells[j].text = "\n".join(parts) + f"\n{total_str}"
                    except (ValueError, TypeError):
                        row_cells[j].text = cell_text
                else:
                    row_cells[j].text = cell_text

        # Legend
        self._add_table_legend(doc, best_ids, second_ids)

    def _add_scales_timeline_chart(
        self, doc: Document, df_session: pd.DataFrame
    ) -> None:
        """Add a rainbow-colored timeline chart of scale trends with a general index line."""
        import math as _math
        from collections import defaultdict
        from io import BytesIO

        doc.add_heading("Scale Trends", level=2)

        # Guard: need valid input
        if df_session is None or df_session.empty:
            doc.add_paragraph("No session data available for chart.")
            return
        if "scale_name" not in df_session.columns or "scale_value" not in df_session.columns:
            doc.add_paragraph("No scale columns found in session data.")
            return

        # Build source-file index (chronological order)
        sources = []
        if "_source_file" in df_session.columns:
            for s in df_session["_source_file"].unique():
                if s not in sources:
                    sources.append(s)
        if not sources:
            doc.add_paragraph("No source files found in session data.")
            return

        source_idx = {s: i for i, s in enumerate(sources)}

        # Collect raw (session_idx, value) pairs per scale
        raw_data = {}  # scale_name -> {session_idx: [values]}
        for _, row in df_session.iterrows():
            sname = str(row.get("scale_name", "") or "").strip()
            sval = str(row.get("scale_value", "") or "").strip()
            src = str(row.get("_source_file", "") or "").strip()
            if not sname or not sval or not src:
                continue
            try:
                val = float(sval)
            except ValueError:
                continue
            if _math.isnan(val):
                continue
            sidx = source_idx.get(src, 0)
            raw_data.setdefault(sname, defaultdict(list))[sidx].append(val)

        if not raw_data:
            doc.add_paragraph("No numeric scale values recorded across sessions.")
            return

        # Average values per session for each scale
        scale_data = {}  # scale_name -> {session_idx: avg_value}
        for name, by_session in raw_data.items():
            scale_data[name] = {si: sum(vs) / len(vs) for si, vs in by_session.items()}

        try:
            import pyqtgraph as pg
            from PySide6.QtCore import QBuffer, QIODevice, Qt
            from PySide6.QtGui import QBrush, QColor, QFont, QPen

            pg.setConfigOptions(useOpenGL=False, antialias=True)

            n_scales = len(scale_data)
            rainbow = [QColor.fromHsvF(i / max(n_scales, 1), 0.85, 0.85)
                        for i in range(n_scales)]

            # Session tick labels
            tick_labels = [s.replace("_events.tsv", "").replace(".tsv", "")
                           for s in sources]
            x_ticks = [(i, lbl) for i, lbl in enumerate(tick_labels)]

            has_index = n_scales >= 2
            win = pg.GraphicsLayoutWidget()
            win.setBackground('w')
            win.resize(1050, 500)  # Single plot, larger for right-side legend

            # --- Main scales chart with General Index on same plot ---
            p1 = win.addPlot(row=0, col=0)
            p1.setTitle("Longitudinal Scale Trends", color='k', size='14pt')
            p1.setLabel('left', 'Scale Value', color='k', size='14pt', font='Arial')
            p1.setLabel('bottom', 'Session', color='k', size='14pt', font='Arial')
            p1.getAxis('bottom').setTicks([x_ticks])
            p1.getAxis('left').setStyle(tickFont=QFont('Arial', 10))
            p1.getAxis('bottom').setStyle(tickFont=QFont('Arial', 10))
            p1.showGrid(x=True, y=True, alpha=0.3)
            # Legend on right side external - increase offset and add background
            legend = p1.addLegend(offset=(1.15, 0.5), pen=QPen(Qt.black, 1), brush=QBrush(Qt.white))
            legend.setLabelTextColor('k')

            # Plot individual scales with original values (no normalization)
            for idx, (sname, pts) in enumerate(scale_data.items()):
                c = rainbow[idx]
                xs = sorted(pts.keys())
                ys = [pts[x] for x in xs]
                p1.plot(xs, ys,
                        pen=pg.mkPen(c, width=2),
                        symbol='o', symbolPen=pg.mkPen(c, width=1),
                        symbolBrush=pg.mkBrush(c), symbolSize=8,
                        name=sname)

            # --- General Index on same plot (if >= 2 scales) ---
            if has_index:
                all_sessions = sorted({s for pts in scale_data.values() for s in pts})

                # Create scale targets dictionary from preferences
                scale_targets = {}
                if self.scale_optimization_prefs:
                    for pref in self.scale_optimization_prefs:
                        if len(pref) >= 5:
                            name, smin, smax, mode, custom_val = pref
                            if mode == "min":
                                scale_targets[name] = {"type": "min", "value": smin}
                            elif mode == "max":
                                scale_targets[name] = {"type": "max", "value": smax}
                            elif mode == "custom":
                                try:
                                    scale_targets[name] = {"type": "custom", "value": float(custom_val)}
                                except ValueError:
                                    scale_targets[name] = {"type": "custom", "value": 0.0}

                index_vals = {}
                for s in all_sessions:
                    weighted_scores = []
                    weights = []

                    for scale_name in scale_data:
                        if s in scale_data[scale_name]:
                            original_value = scale_data[scale_name][s]

                            # Get target for this scale
                            if scale_name in scale_targets:
                                target_info = scale_targets[scale_name]
                                target_type = target_info["type"]
                                target_value = target_info["value"]

                                # Calculate distance from target (lower is better)
                                if target_type == "min":
                                    # For minimization: lower values are better
                                    distance = original_value
                                    max_possible = max(scale_data[scale_name].values())
                                    # Normalize: 0 = at target (min), 1 = worst (max)
                                    normalized_score = distance / max_possible if max_possible > 0 else 0
                                elif target_type == "max":
                                    # For maximization: higher values are better
                                    distance = max(scale_data[scale_name].values()) - original_value
                                    max_possible = max(scale_data[scale_name].values()) - min(scale_data[scale_name].values())
                                    # Normalize: 0 = at target (max), 1 = worst (min)
                                    normalized_score = distance / max_possible if max_possible > 0 else 0
                                elif target_type == "custom":
                                    # For custom target: absolute distance from target
                                    distance = abs(original_value - target_value)
                                    max_distance = max(abs(v - target_value) for v in scale_data[scale_name].values())
                                    # Normalize: 0 = at target, 1 = worst
                                    normalized_score = distance / max_distance if max_distance > 0 else 0

                                # Convert to proximity score (higher is better)
                                proximity_score = 1.0 - normalized_score
                                weighted_scores.append(proximity_score)
                                weights.append(1.0)  # Equal weight for now
                            else:
                                # No target defined: use neutral score
                                weighted_scores.append(0.5)  # Neutral middle value
                                weights.append(0.5)  # Lower weight for scales without targets

                    if weighted_scores and weights:
                        # Calculate weighted average of proximity scores
                        total_weight = sum(weights)
                        if total_weight > 0:
                            index_vals[s] = sum(w * s for w, s in zip(weights, weighted_scores, strict=False)) / total_weight
                        else:
                            index_vals[s] = 0.5  # Default neutral value

                if index_vals:
                    ix = sorted(index_vals.keys())
                    iy = [index_vals[x] for x in ix]
                    # Thicker black line for General Index
                    p1.plot(ix, iy,
                            pen=pg.mkPen('k', width=5),
                            symbol='d', symbolPen='k', symbolBrush='k',
                            symbolSize=10, name='General Index')

            # --- Export to PNG → Word ---
            pixmap = win.grab()
            qbuf = QBuffer()
            qbuf.open(QIODevice.OpenModeFlag.WriteOnly)
            pixmap.save(qbuf, 'PNG')
            qbuf.close()
            img_buf = BytesIO(bytes(qbuf.data()))
            doc.add_picture(img_buf, width=Inches(6))
            doc.add_paragraph()
            img_buf.close()
            win.close()
            del win

        except Exception as e:
            doc.add_paragraph(f"Chart generation error: {e}")

    # ------------------------------------------------------------------
    # Data helpers
    # ------------------------------------------------------------------

    def _create_lateral_table(self, df: pd.DataFrame) -> pd.DataFrame:
        """Create lateral (L/R) table structure similar to SessionExporter."""
        if df.empty:
            return df

        df = self._normalize_block_id(df)

        # Create a global entry id combining source file + block_id
        if "_source_file" in df.columns and "block_id" in df.columns:
            df["_global_entry_id"] = df["_source_file"] + "_" + df["block_id"].astype(str)
        elif "block_id" in df.columns:
            df["_global_entry_id"] = df["block_id"].astype(str)
        else:
            df["_global_entry_id"] = range(len(df))

        groups = df.groupby("_global_entry_id", sort=False, dropna=False)
        lateral_data = []

        for entry_id, block_df in groups:
            first = block_df.iloc[0]

            # Collect scales (filter out NaN values)
            scale_pairs = []
            seen = set()
            for _, r in block_df.iterrows():
                sn = str(r.get("scale_name", "") or "").strip()
                sv = str(r.get("scale_value", "") or "").strip()

                # Skip if scale value is NaN or empty
                if not sv or sv == "NaN" or sv.strip() == "NaN":
                    continue

                if sn and (sn, sv) not in seen:
                    seen.add((sn, sv))
                    scale_pairs.append((sn, sv))

            combined_sn = "\n".join(p[0] for p in scale_pairs) if scale_pairs else ""
            combined_sv = "\n".join(p[1] for p in scale_pairs) if scale_pairs else ""

            source_label = str(first.get("_source_file", "")).replace("_events.tsv", "").replace(".tsv", "")
            date_val = str(first.get("date", "") or "")

            common = {
                "_global_entry_id": entry_id,
                "source": source_label,
                "date": date_val,
                "program_ID": first.get("program_ID", ""),
                "scale_name": combined_sn,
                "scale_value": combined_sv,
                "notes": first.get("notes", ""),
            }

            lat_map = {
                "left_stim_freq": "frequency", "left_cathode": "cathode",
                "left_anode": "anode", "left_amplitude": "amplitude",
                "left_pulse_width": "pulse_width",
                "right_stim_freq": "frequency", "right_cathode": "cathode",
                "right_anode": "anode", "right_amplitude": "amplitude",
                "right_pulse_width": "pulse_width",
            }

            left = dict(common)
            right = dict(common)
            left["laterality"] = "L"
            right["laterality"] = "R"

            for col, generic in lat_map.items():
                if col.startswith("left_"):
                    left[generic] = first.get(col, "")
                else:
                    right[generic] = first.get(col, "")

            lateral_data.append(left)
            lateral_data.append(right)

        return pd.DataFrame(lateral_data)

    def _find_best_and_second_best(self, lateral_df: pd.DataFrame) -> tuple:
        """Find entry IDs with the best and second-best scores."""
        if lateral_df is None or lateral_df.empty:
            return [], []
        if "_global_entry_id" not in lateral_df.columns:
            return [], []
        if "scale_name" not in lateral_df.columns or "scale_value" not in lateral_df.columns:
            return [], []

        try:
            pref_lookup = {}
            for pref in self.scale_optimization_prefs:
                if len(pref) >= 5:
                    name, _, _, mode, custom_val = pref
                    pref_lookup[name.strip().lower()] = (mode, custom_val)

            df_l = lateral_df[lateral_df.get("laterality", "") == "L"].copy()
            if df_l.empty:
                df_l = lateral_df.drop_duplicates(subset=["_global_entry_id"]).copy()

            scores = {}
            for _, row in df_l.iterrows():
                eid = row.get("_global_entry_id")
                if eid is None:
                    continue
                names = str(row.get("scale_name", "") or "").split("\n")
                values = str(row.get("scale_value", "") or "").split("\n")

                total = 0.0
                has_val = False
                import math as _math
                for i, vl in enumerate(values):
                    vl = vl.strip()
                    if not vl:
                        continue
                    try:
                        val = float(vl)
                    except ValueError:
                        continue
                    if _math.isnan(val):
                        continue
                    sn = names[i].strip().lower() if i < len(names) else ""
                    mode, cv = pref_lookup.get(sn, ("min", ""))
                    if mode == "ignore":
                        continue
                    has_val = True
                    if mode in ("low", "min"):
                        total += val
                    elif mode in ("high", "max"):
                        total -= val
                    elif mode == "custom":
                        try:
                            total += abs(val - float(cv))
                        except ValueError:
                            total += val
                if has_val:
                    scores[eid] = total

            if not scores:
                return [], []

            unique = sorted(set(scores.values()))
            best = [eid for eid, s in scores.items() if s == unique[0]]
            second = (
                [eid for eid, s in scores.items() if s == unique[1]]
                if len(unique) > 1
                else []
            )
            return best, second
        except Exception:
            return [], []

    # ------------------------------------------------------------------
    # Formatting / utility helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_manufacturer_for_model(model_name: str) -> str:
        """Return the manufacturer string for a given electrode model name."""
        if not model_name:
            return ""
        for manufacturer, models in (MANUFACTURERS or {}).items():
            try:
                if model_name in models:
                    return str(manufacturer)
            except Exception:
                continue
        return ""

    def _render_electrode_png(
        self,
        model_name: str,
        anode_text: str,
        cathode_text: str,
        target_size_px: tuple = (440, 900),
    ) -> str | None:
        """Render electrode configuration to a temporary PNG file."""
        try:
            from PySide6.QtGui import QColor as _QColor
            from PySide6.QtGui import QPainter, QPixmap

            from ..models import ElectrodeCanvas

            model = ELECTRODE_MODELS.get(model_name)
            if not model:
                return None

            canvas = ElectrodeCanvas()
            canvas.set_model(model)
            canvas.resize(*target_size_px)
            try:
                canvas.set_export_mode(True)
            except Exception:
                pass

            # Apply contact states
            canvas.contact_states.clear()
            canvas.case_state = ContactState.OFF

            def apply_tokens(text: str, state: ContactState) -> None:
                if not text:
                    return
                for token in str(text).split("_"):
                    token = token.strip()
                    if not token:
                        continue
                    if token == "case":
                        canvas.case_state = state
                        continue
                    if token.startswith("E") and len(token) >= 2:
                        try:
                            if token[-1].isalpha():
                                idx = int(token[1:-1])
                                seg_map = {"a": 0, "b": 1, "c": 2}
                                seg_char = token[-1].lower()
                                if seg_char in seg_map:
                                    canvas.contact_states[(idx, seg_map[seg_char])] = state
                            else:
                                idx = int(token[1:])
                                if model.is_directional:
                                    for seg in range(3):
                                        canvas.contact_states[(idx, seg)] = state
                                else:
                                    canvas.contact_states[(idx, 0)] = state
                        except Exception:
                            continue

            apply_tokens(anode_text, ContactState.ANODIC)
            apply_tokens(cathode_text, ContactState.CATHODIC)
            canvas.update()

            # Render with white background
            original_paint = canvas.paintEvent
            def white_bg_paint(event):
                painter = QPainter(canvas)
                painter.fillRect(canvas.rect(), Qt.white)
                original_paint(event)
            canvas.paintEvent = white_bg_paint

            pixmap = QPixmap(canvas.size())
            pixmap.fill(Qt.white)
            canvas.render(pixmap)

            # Crop white borders
            image = pixmap.toImage()
            white_rgb = _QColor(Qt.white).rgb()
            left, top, right, bottom = image.width(), image.height(), 0, 0
            for y in range(image.height()):
                for x in range(image.width()):
                    if image.pixel(x, y) != white_rgb:
                        left = min(left, x)
                        top = min(top, y)
                        right = max(right, x)
                        bottom = max(bottom, y)
            if right > left and bottom > top:
                margin = 20
                left = max(0, left - margin)
                top = max(0, top - margin)
                right = min(image.width() - 1, right + margin)
                bottom = min(image.height() - 1, bottom + margin)
                cropped = pixmap.copy(left, top, right - left + 1, bottom - top + 1)
            else:
                cropped = pixmap

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.close()
            cropped.save(tmp.name, "PNG")
            return tmp.name
        except Exception:
            return None

    @staticmethod
    def _pick_latest_row(df: pd.DataFrame):
        """Return the row with the highest block_id, or the last row if unavailable."""
        if df is None or df.empty:
            return None
        for col in ("block_id", "block_ID", "blockId", "blockID"):
            if col in df.columns:
                try:
                    numeric = pd.to_numeric(df[col], errors="coerce")
                    max_val = numeric.max()
                    if pd.notna(max_val):
                        return df.loc[numeric == max_val].iloc[-1]
                except Exception:
                    pass
        return df.iloc[-1]

    @staticmethod
    def _normalize_block_id(df: pd.DataFrame) -> pd.DataFrame:
        if df is None or df.empty or "block_id" in df.columns:
            return df
        for c in ("block_ID", "blockId", "blockID"):
            if c in df.columns:
                return df.rename(columns={c: "block_id"})
        return df

    @staticmethod
    def _column_header(col: str) -> str:
        m = {
            "source": "#",
            "date": "Date",
            "scale_name": PLACEHOLDERS.get("scale_name", "Scale"),
            "scale_value": PLACEHOLDERS.get("scale_value", "Value"),
            "frequency": PLACEHOLDERS.get("frequency", "Freq"),
            "anode": "+", "cathode": "-",
            "amplitude": PLACEHOLDERS.get("amplitude", "Amp"),
            "pulse_width": PLACEHOLDERS.get("pulse_width", "PW"),
            "program_ID": "Prog", "laterality": "",
        }
        return m.get(col, col.replace("_", " ").title())

    @staticmethod
    def _extract_patient_id(file_paths: list[str]) -> str:
        for fp in file_paths:
            m = re.search(r"sub-([^_]+)", os.path.basename(fp))
            if m:
                return m.group(1)
        return ""

    @staticmethod
    def _generate_filename(file_paths: list[str], ext: str) -> str:
        today = datetime.now().astimezone().strftime("%Y%m%d")
        pid = LongitudinalExporter._extract_patient_id(file_paths)
        if pid:
            return f"sub-{pid}_longitudinal-report_{today}{ext}"
        return f"longitudinal-report_{today}{ext}"

    def _highlight_cells(self, row_cells, intensity: str = "best") -> None:
        color = "C3E6CB" if intensity == "best" else "E8F5E9"
        for cell in row_cells:
            try:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), color)
                cell._tc.get_or_add_tcPr().append(shd)
            except Exception:
                pass

    @staticmethod
    def _set_cell_border_top(cell, sz=12):
        try:
            tcPr = cell._tc.get_or_add_tcPr() # noqa: N806
            borders = OxmlElement("w:tcBorders")
            top = OxmlElement("w:top")
            top.set(qn("w:val"), "single")
            top.set(qn("w:sz"), str(sz))
            top.set(qn("w:space"), "0")
            top.set(qn("w:color"), "000000")
            borders.append(top)
            tcPr.append(borders)
        except Exception:
            pass

    def _add_table_legend(self, doc: Document, best_ids: list, second_ids: list) -> None:
        if not best_ids and not second_ids:
            return

        doc.add_paragraph()

        legend = doc.add_paragraph()
        legend.add_run("Legend: ").bold = True
        if best_ids:
            r = legend.add_run("■ ")
            r.font.color.rgb = RGBColor(0xC3, 0xE6, 0xCB)
            legend.add_run("Optimal entry    ")
        if second_ids:
            r = legend.add_run("■ ")
            r.font.color.rgb = RGBColor(0xE8, 0xF5, 0xE9)
            legend.add_run("Second-best entry")

        if self.scale_optimization_prefs:
            tp = doc.add_paragraph()
            tp.add_run("Scale targets: ").bold = True
            parts = []
            for pref in self.scale_optimization_prefs:
                if len(pref) >= 5:
                    name, smin, smax, mode, cv = pref
                    if mode == "ignore":
                        continue
                    elif mode == "min":
                        parts.append(f"{name}: min")
                    elif mode == "max":
                        parts.append(f"{name}: max")
                    elif mode == "custom":
                        parts.append(f"{name}: {cv}")
            if parts:
                tp.add_run("; ".join(parts))
                for run in tp.runs:
                    run.font.size = Pt(9)

        disc = doc.add_paragraph()
        dr = disc.add_run(
            "Note: The highlighted rows are derived exclusively from the recorded "
            "session scale values and represent a computational ranking intended "
            "solely as a reference. This color-coded indication does not constitute "
            "clinical guidance."
        )
        dr.font.size = Pt(9)
        dr.font.italic = True

    def _show_transient_message(
        self, parent, title: str, text: str, msecs: int = 2000
    ) -> None:
        msg = QMessageBox(parent)
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle(title)
        msg.setText(text)
        msg.setStandardButtons(QMessageBox.NoButton)
        msg.setWindowModality(Qt.NonModal)
        msg.show()

        timer = QTimer(msg)
        timer.setSingleShot(True)

        def _close():
            try:
                msg.accept()
            except Exception:
                pass

        timer.timeout.connect(_close)
        timer.start(max(0, int(msecs)))

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """Convert Word → PDF using the same strategy as SessionExporter."""
        import shutil
        import subprocess

        errors = []

        try:
            from docx2pdf import convert as _convert
            _convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return
        except Exception as e:
            errors.append(f"docx2pdf: {e}")

        try:
            abs_d = os.path.abspath(docx_path).replace("'", "''")
            abs_p = os.path.abspath(pdf_path).replace("'", "''")
            ps = (
                "$w = New-Object -ComObject Word.Application; "
                "$w.Visible = $false; "
                f"$d = $w.Documents.Open('{abs_d}'); "
                f"$d.SaveAs2('{abs_p}', 17); "
                "$d.Close(); $w.Quit()"
            )
            subprocess.run(
                ["powershell", "-NoProfile", "-Command", ps],
                check=True, capture_output=True, timeout=60,
            )
            if os.path.exists(pdf_path):
                return
        except Exception as e:
            errors.append(f"Word COM: {e}")

        soffice = shutil.which("soffice")
        if soffice:
            try:
                out_dir = os.path.dirname(os.path.abspath(pdf_path))
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", out_dir, os.path.abspath(docx_path)],
                    check=True, capture_output=True, timeout=60,
                )
                lo_out = os.path.join(
                    out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
                )
                if lo_out != pdf_path and os.path.exists(lo_out):
                    shutil.move(lo_out, pdf_path)
                if os.path.exists(pdf_path):
                    return
            except Exception as e:
                errors.append(f"LibreOffice: {e}")

        raise RuntimeError(
            "Could not convert to PDF:\n" + "\n".join(errors)
            + "\n\nPlease export to Word and convert manually."
        )
