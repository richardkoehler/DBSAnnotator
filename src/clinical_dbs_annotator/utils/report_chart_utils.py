"""
Shared chart utilities for report generation.

Provides reusable functions for building pyqtgraph-based scale timeline
charts used by both the session and longitudinal exporters.  Centralising
the chart logic avoids code duplication and guarantees consistent styling
across report types.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx.document import Document
from docx.shared import Inches
from matplotlib import cm
from pyqtgraph import AxisItem
from PySide6.QtCore import QRectF
from PySide6.QtGui import QPainter

# ── Colour constants ────────────────────────────────────────────────
BEST_GREEN = (150, 210, 160, 160)  # RGBA – darker green for best config
SECOND_GREEN = (200, 235, 205, 130)  # RGBA – lighter green for second-best


# ── Scale-target helpers ────────────────────────────────────────────


def parse_scale_targets(
    prefs: list[tuple[str, str, str, str, str]],
) -> dict[str, dict[str, Any]]:
    """Convert user preference tuples into a look-up dict.

    Args:
        prefs: [(name, min, max, mode, custom_value), ...]

    Returns:
        {scale_name: {"type": mode, "value": numeric_target}}
    """
    targets: dict[str, dict[str, Any]] = {}
    for pref in prefs or []:
        if len(pref) < 5:
            continue
        name, smin, smax, mode, custom_val = pref
        if mode == "min":
            targets[name] = {"type": "min", "value": float(smin) if smin else 0.0}
        elif mode == "max":
            targets[name] = {"type": "max", "value": float(smax) if smax else 0.0}
        elif mode == "custom":
            try:
                targets[name] = {"type": "custom", "value": float(custom_val)}
            except ValueError, TypeError:
                targets[name] = {"type": "custom", "value": 0.0}
    return targets


def compute_general_index(
    scale_data: dict[str, dict[int, float]],
    all_points: list[int],
    scale_targets: dict[str, dict[str, Any]],
) -> dict[int, float]:
    """Compute a weighted general-index value per x-point.

    The index represents how close the scales are to their respective
    targets (1.0 = perfect, 0.0 = worst observed).

    Args:
        scale_data:    {scale_name: {x_point: value}}
        all_points:    sorted list of x-axis keys
        scale_targets: output of :func:`parse_scale_targets`

    Returns:
        {x_point: index_value}
    """
    index_vals: dict[int, float] = {}
    for pt in all_points:
        weighted_scores: list[float] = []
        weights: list[float] = []

        for scale_name, pts in scale_data.items():
            if pt not in pts:
                continue
            original_value = pts[pt]

            if scale_name in scale_targets:
                info = scale_targets[scale_name]
                ttype = info["type"]
                tvalue = info["value"]

                if ttype == "min":
                    distance = original_value
                    max_possible = max(pts.values())
                    normalized = distance / max_possible if max_possible > 0 else 0
                elif ttype == "max":
                    distance = max(pts.values()) - original_value
                    max_possible = max(pts.values()) - min(pts.values())
                    normalized = distance / max_possible if max_possible > 0 else 0
                elif ttype == "custom":
                    distance = abs(original_value - tvalue)
                    max_distance = max(abs(v - tvalue) for v in pts.values())
                    normalized = distance / max_distance if max_distance > 0 else 0
                else:
                    normalized = 0.5

                weighted_scores.append(1.0 - normalized)
                weights.append(1.0)
            else:
                weighted_scores.append(0.5)
                weights.append(0.5)

        if weighted_scores and weights:
            total_w = sum(weights)
            if total_w > 0:
                index_vals[pt] = (
                    sum(w * s for w, s in zip(weights, weighted_scores, strict=False))
                    / total_w
                )
            else:
                index_vals[pt] = 0.5

    return index_vals


def find_best_and_second(
    index_vals: dict[int, float],
) -> tuple[int | None, int | None]:
    """Return the x-points with the best and second-best general-index scores.

    Returns:
        (best_x, second_best_x) – either may be *None*.
    """
    if not index_vals:
        return None, None
    ranked = sorted(index_vals, key=lambda k: index_vals[k], reverse=True)
    best = ranked[0]
    second = ranked[1] if len(ranked) > 1 else None
    return best, second


# ── Chart rendering ────────────────────────────────────────────────


class RotatedAxisItem(AxisItem):
    def __init__(self, orientation, angle=0, **kwargs):
        super().__init__(orientation, **kwargs)
        self.angle = angle

    def drawPicture(self, p: QPainter, axisSpec, tickSpecs, textSpecs):  # noqa: N802, N803
        super().drawPicture(p, axisSpec, tickSpecs, [])

        # Find maximum label width for consistent vertical offset
        max_width = max((rect.width() for rect, _, _ in textSpecs), default=0)

        # Rotate ticks labels
        for rect, flags, text in textSpecs:
            p.save()
            # Translate to tick mark position (bottom-left of original rect), then rotate counter-clockwise
            p.translate(rect.bottomLeft())
            p.rotate(self.angle)
            # Draw text extending downward, offset by max_width for consistent alignment
            new_rect = QRectF(0, -max_width, rect.width(), rect.height())
            p.drawText(new_rect, flags, text)
            p.restore()


def build_scales_chart(
    scale_data: dict[str, dict[int, float]],
    scale_prefs: list[tuple[str, str, str, str, str]] | None,
    *,
    title: str = "",
    x_label: str = "X",
    y_label: str = "Scale Value",
    x_ticks: list[tuple[int, str]] | None = None,
    width: int = 1100,
    height: int = 520,
    show_general_index: bool = True,
    rotate_x_ticks: bool = False,
) -> bytes | None:
    """Build a scale-trend chart and return it as PNG bytes.

    The chart includes:
    * Rainbow-coloured individual scale lines
    * A thick black General Index line (when *show_general_index* and ≥ 2 scales)
    * Green vertical bands for best (dark) and second-best (light)
    * A compact legend anchored at the **top-right** of the plot

    Args:
        scale_data:          {scale_name: {x_point: value}}
        scale_prefs:         user optimisation preferences (may be *None*)
        title:               chart title text (empty string for no title)
        x_label:             bottom-axis label
        y_label:             left-axis label
        x_ticks:             optional custom bottom-axis ticks
        width:               image width in px
        height:              image height in px
        show_general_index:  whether to draw the General Index line
        rotate_x_ticks:      whether to rotate x-axis tick labels by 70 degrees

    Returns:
        PNG image bytes, or *None* on failure.
    """
    try:
        import pyqtgraph as pg
        from PySide6.QtCore import QBuffer, QIODevice, Qt
        from PySide6.QtGui import QBrush, QColor, QFont, QPen

        pg.setConfigOptions(useOpenGL=False, antialias=True)

        n_scales = len(scale_data)
        if n_scales == 0:
            return None

        # Get colors from matplotlib dark2 colormap
        cmap = cm.get_cmap("Dark2")
        colors = []
        for i in range(n_scales):
            rgba = cmap(i % cmap.N)
            colors.append(QColor.fromRgbF(*rgba[:3]))

        # Line styles for each curve
        line_styles = [
            Qt.SolidLine,
            Qt.DashLine,
            Qt.DotLine,
            Qt.DashDotLine,
            Qt.DashDotDotLine,
        ]

        has_index = show_general_index and n_scales >= 2
        total_items = n_scales + (1 if has_index else 0)

        win = pg.GraphicsLayoutWidget()
        win.setBackground("w")
        # Increase height when rotating ticks to compensate for larger bottom margin
        actual_height = height + 150 if rotate_x_ticks else height
        win.resize(width, actual_height)

        # Legend in a dedicated row above the plot (compact)
        legend = pg.LegendItem(
            offset=(0, 0),
            colCount=max(1, total_items),
        )
        legend.setBrush(QBrush(QColor(255, 255, 255, 220)))
        legend.setPen(QPen(QColor(150, 150, 150, 100), 1))
        legend.setLabelTextColor("k")
        win.addItem(legend, row=0, col=0)
        win.ci.layout.setRowStretchFactor(0, 0)  # legend row: no stretch
        win.ci.layout.setRowStretchFactor(1, 1)  # plot row: takes space

        p1 = win.addPlot(row=1, col=0)

        if rotate_x_ticks:
            axis = RotatedAxisItem(orientation="bottom", angle=-90)
            axis.linkToView(p1.getViewBox())
            p1.setAxisItems({"bottom": axis})
            p1.layout.setContentsMargins(10, 10, 10, 200)

        if title:
            p1.setTitle(title, color="k", size="16pt", font="Arial")
        p1.setLabel("left", y_label, color="k", **{"font-size": "16pt"})
        p1.setLabel("bottom", x_label, color="k", **{"font-size": "16pt"})
        p1.getAxis("left").setStyle(tickFont=QFont("Arial", 10))
        p1.getAxis("bottom").setStyle(tickFont=QFont("Arial", 10))

        if x_ticks is not None:
            p1.getAxis("bottom").setTicks([x_ticks])
        p1.showGrid(x=True, y=True, alpha=0.3)

        # ── Plot individual scales ──────────────────────────────────
        for idx, (sname, pts) in enumerate(scale_data.items()):
            c = colors[idx]
            ls = line_styles[idx % len(line_styles)]
            xs = sorted(pts.keys())
            ys = [pts[x] for x in xs]
            curve = p1.plot(
                xs,
                ys,
                pen=pg.mkPen(c, width=2, style=ls),
                symbol="o",
                symbolPen=pg.mkPen(c, width=1),
                symbolBrush=pg.mkBrush(c),
                symbolSize=8,
            )
            legend.addItem(curve, sname)

        # ── General Index ───────────────────────────────────────────
        best_x: int | None = None
        second_x: int | None = None
        if has_index:
            all_points = sorted({x for pts in scale_data.values() for x in pts})
            scale_targets = parse_scale_targets(scale_prefs)
            index_vals = compute_general_index(scale_data, all_points, scale_targets)

            if index_vals:
                ix = sorted(index_vals.keys())
                iy = [index_vals[x] for x in ix]
                gi_curve = p1.plot(
                    ix,
                    iy,
                    pen=pg.mkPen("k", width=5),
                    symbol="d",
                    symbolPen="k",
                    symbolBrush="k",
                    symbolSize=10,
                )
                legend.addItem(gi_curve, "General Index")
                best_x, second_x = find_best_and_second(index_vals)

        # ── Best / second-best green vertical bands ─────────────────
        def _add_band(x_val: int, rgba: tuple) -> None:
            band = pg.LinearRegionItem(
                values=(x_val - 0.35, x_val + 0.35),
                orientation="vertical",
                brush=QBrush(QColor(*rgba)),
                pen=pg.mkPen(QColor(100, 180, 100, 120), width=1),
                movable=False,
            )
            band.setZValue(-10)
            p1.addItem(band)

        if best_x is not None:
            _add_band(best_x, BEST_GREEN)
        if second_x is not None and second_x != best_x:
            _add_band(second_x, SECOND_GREEN)

        # ── Export to PNG bytes ─────────────────────────────────────
        pixmap = win.grab()
        qbuf = QBuffer()
        qbuf.open(QIODevice.OpenModeFlag.WriteOnly)
        pixmap.save(qbuf, "PNG")
        qbuf.close()
        png_bytes = bytes(qbuf.data())

        win.close()
        del win
        return png_bytes

    except Exception as exc:
        import logging

        logging.getLogger(__name__).exception("build_scales_chart failed: %s", exc)
        return None


def add_chart_to_doc(
    doc: Document,
    png_bytes: bytes | None,
    *,
    heading: str | None = None,
    heading_level: int = 2,
    width_inches: float | None = None,
    fallback_message: str = "Chart generation error.",
) -> None:
    """Insert a PNG chart into a Word document.

    Args:
        doc:              python-docx Document instance
        png_bytes:        raw PNG bytes (or *None* on failure)
        heading:          optional heading text above the chart
        heading_level:    heading level (default 2)
        width_inches:     image width in the document (None for full page width)
        fallback_message: text shown if *png_bytes* is None
    """
    if heading:
        doc.add_heading(heading, level=heading_level)

    # Calculate page width if not specified
    if width_inches is None:
        section = doc.sections[0]
        page_w = (
            section.page_width - section.left_margin - section.right_margin
        ) / 914400  # Convert twips to inches
        width_inches = max(4.0, page_w)  # Minimum 4 inches, otherwise full page

    if png_bytes is None:
        doc.add_paragraph(fallback_message)
        return

    img_buf = BytesIO(png_bytes)
    doc.add_picture(img_buf, width=Inches(width_inches))
    doc.add_paragraph()
    img_buf.close()
