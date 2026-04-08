"""
Session data exporter for Clinical DBS Annotator.

This module provides functionality to export session data to Word and PDF.
"""

import csv
import os
import shutil
import subprocess
import tempfile
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PySide6.QtCore import Qt, QTimer
from PySide6.QtGui import QPainter, QPixmap
from PySide6.QtWidgets import QMessageBox, QWidget

from .. import __app_name__, __version__
from ..config import PLACEHOLDERS
from ..config_electrode_models import ELECTRODE_MODELS, MANUFACTURERS, ContactState
from ..models import ElectrodeCanvas


class SessionExporter:
    """
    Handles exporting session data to various formats.

    This class provides methods to export the collected session data
    to Word and PDF.
    """

    def __init__(self, session_data):
        """
        Initialize the session exporter.

        Args:
            session_data: The SessionData instance containing collected data
        """
        self.session_data = session_data
        # Scale optimization preferences: list of (name, min, max, mode, custom_value)
        # mode: "low", "high", "custom", "ignore"
        self.scale_optimization_prefs: list = []

    def set_scale_optimization_prefs(self, prefs: list) -> None:
        """Set scale optimization preferences for best block calculation."""
        self.scale_optimization_prefs = prefs or []

    def _generate_bids_report_filename(self, extension: str = '.docx') -> str:
        """Generate BIDS-friendly report filename from TSV file path."""
        tsv_path = getattr(self.session_data, 'file_path', '') or ''
        today_str = datetime.now().astimezone().strftime('%Y%m%d')

        if tsv_path:
            base = os.path.basename(tsv_path)
            import re
            sub_match = re.search(r'sub-([^_]+)', base)
            run_match = re.search(r'run-([0-9]+)', base)
            task_match = re.search(r'task-([^_]+)', base)
            patient_id = sub_match.group(1) if sub_match else "unknown"
            run_num = run_match.group(1) if run_match else "01"
            task = task_match.group(1) if task_match else "programming"

            return f"sub-{patient_id}_ses-{today_str}_task-{task}_run-{run_num}_report{extension}"

        # Fallback
        return f"dbs_session_report_{today_str}_{datetime.now().astimezone().strftime('%H%M%S')}{extension}"

    def _extract_bids_info_from_path(self) -> tuple:
        """Extract patient ID and session number from BIDS filename."""
        tsv_path = getattr(self.session_data, 'file_path', '') or ''
        patient_id = ""
        session_num = ""
        if tsv_path:
            base = os.path.basename(tsv_path)
            # Parse sub-XXX and ses-XXX from filename
            import re
            sub_match = re.search(r'sub-([^_]+)', base)
            ses_match = re.search(r'ses-([^_]+)', base)
            if sub_match:
                patient_id = sub_match.group(1)
            if ses_match:
                raw_session = ses_match.group(1)
                # Format session date from 20250103 to 2025-01-03
                if len(raw_session) == 8 and raw_session.isdigit():
                    try:
                        year = raw_session[:4]
                        month = raw_session[4:6]
                        day = raw_session[6:8]
                        session_num = f"{year}-{month}-{day}"
                    except Exception:
                        session_num = raw_session
                else:
                    session_num = raw_session
        return patient_id, session_num

    def _show_transient_message(
        self,
        parent: QWidget | None,
        title: str,
        text: str,
        *,
        msecs: int = 2000,
        icon: QMessageBox.Icon = QMessageBox.Information,
    ) -> None:
        msg = QMessageBox(parent)
        msg.setIcon(icon)
        msg.setWindowTitle(title)
        msg.setText(text)
        msg.setStandardButtons(QMessageBox.NoButton)
        msg.setWindowModality(Qt.NonModal)
        msg.show()

        if parent is not None:
            try:
                parent._export_transient_msg = msg
            except Exception:
                pass

        # Use a dedicated QTimer owned by the message box so it reliably fires.
        # Some Qt builds do not ship QWeakPointer.
        timer = QTimer(msg)
        timer.setSingleShot(True)

        def _close_msg() -> None:
            try:
                msg.accept()
            except Exception:
                try:
                    msg.close()
                except Exception:
                    pass

            if parent is not None:
                try:
                    if getattr(parent, "_export_transient_msg", None) is msg:
                        parent._export_transient_msg = None
                except Exception:
                    pass

        timer.timeout.connect(_close_msg)
        timer.start(max(0, int(msecs)))


    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """Convert a Word document to PDF using the best available method.

        Tries in order:
        1. docx2pdf (requires Microsoft Word COM)
        2. Word COM via PowerShell subprocess
        3. LibreOffice headless

        Raises RuntimeError if no conversion method succeeds.
        """
        errors: list[str] = []

        # 1. Try docx2pdf
        try:
            from docx2pdf import convert as _docx2pdf_convert
            _docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return
        except Exception as exc:
            errors.append(f"docx2pdf: {exc}")

        # 2. Try Word COM via PowerShell (Windows)
        try:
            abs_docx = os.path.abspath(docx_path).replace("'", "''")
            abs_pdf = os.path.abspath(pdf_path).replace("'", "''")
            ps_script = (
                "$w = New-Object -ComObject Word.Application; "
                "$w.Visible = $false; "
                f"$d = $w.Documents.Open('{abs_docx}'); "
                f"$d.SaveAs2('{abs_pdf}', 17); "
                "$d.Close(); $w.Quit()"
            )
            subprocess.run(
                ["powershell", "-NoProfile", "-Command", ps_script],
                check=True, capture_output=True, timeout=60,
            )
            if os.path.exists(pdf_path):
                return
        except Exception as exc:
            errors.append(f"Word COM (PowerShell): {exc}")

        # 3. Try LibreOffice headless
        soffice = shutil.which("soffice")
        if soffice:
            try:
                out_dir = os.path.dirname(os.path.abspath(pdf_path))
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", out_dir, os.path.abspath(docx_path)],
                    check=True, capture_output=True, timeout=60,
                )
                # LibreOffice outputs with same basename
                lo_output = os.path.join(
                    out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
                )
                if lo_output != pdf_path and os.path.exists(lo_output):
                    shutil.move(lo_output, pdf_path)
                if os.path.exists(pdf_path):
                    return
            except Exception as exc:
                errors.append(f"LibreOffice: {exc}")
        else:
            errors.append("LibreOffice: soffice not found on PATH")

        detail = "\n".join(errors)
        raise RuntimeError(
            f"Could not convert to PDF. Tried all available methods:\n{detail}\n\n"
            "Please export to Word (.docx) and convert to PDF manually."
        )

    def _read_session_data(self) -> pd.DataFrame | None:
        """
        Read session data from the TSV file.

        Returns:
            DataFrame with session data or None if error
        """
        try:
            if hasattr(self.session_data, 'file_path') and self.session_data.file_path:
                return pd.read_csv(self.session_data.file_path, sep='\t')
            return None
        except Exception:
            return None

    def _normalize_block_id_column(self, df: pd.DataFrame) -> pd.DataFrame:
        """Rename variant block-id column names to the canonical 'block_id'."""
        if df is None or df.empty:
            return df

        if "block_id" in df.columns:
            return df

        for candidate in ("block_ID", "blockId", "blockID"):
            if candidate in df.columns:
                return df.rename(columns={candidate: "block_id"})

        return df

    def _get_manufacturer_for_model(self, model_name: str) -> str:
        """Return the manufacturer string for a given electrode model name."""
        if not model_name:
            return ""
        for manufacturer, models in (MANUFACTURERS or {}).items():
            try:
                if model_name in models:
                    return str(manufacturer)
            except Exception:
                continue
        return ""

    def _pick_latest_session_row(self, df: pd.DataFrame) -> pd.Series | None:
        """Return the row with the highest session_ID and block_id."""
        if df is None or df.empty:
            return None

        # Prefer: latest session_ID, then latest block_id
        if "session_ID" in df.columns:
            try:
                s = pd.to_numeric(df["session_ID"], errors="coerce")
                max_sid = s.max()
                if pd.notna(max_sid):
                    df_sid = df.loc[s == max_sid]
                    return self._pick_latest_row(df_sid)
            except Exception:
                pass

        return self._pick_latest_row(df)

    def _add_summary_section(self, doc: Document, df: pd.DataFrame, df_initial: pd.DataFrame, df_final: pd.DataFrame) -> None:
        """Add the initial clinical notes section to the Word document."""
        doc.add_heading('Initial Clinical Notes', level=1)

        latest_init = self._pick_latest_session_row(df_initial)

        if latest_init is not None:
            init_sid = latest_init.get("session_ID", None)
            if init_sid is not None and "session_ID" in df_initial.columns:
                df_init_session = df_initial[df_initial["session_ID"] == init_sid]
            else:
                df_init_session = df_initial

            scale_names = []
            scale_values = []
            seen = set()
            for _, r in df_init_session.iterrows():
                sname = r.get("scale_name", "")
                sval = r.get("scale_value", "")
                if pd.notna(sname) and str(sname).strip():
                    name_str = str(sname).strip()
                    val_str = str(sval) if pd.notna(sval) else ""
                    key = (name_str, val_str)
                    if key not in seen:
                        seen.add(key)
                        scale_names.append(name_str)
                        scale_values.append(val_str)

            for sn, sv in zip(scale_names, scale_values, strict=False):
                doc.add_paragraph(f'{sn}: {sv}')

            notes = str(latest_init.get('notes', '') or '')
            if notes.strip():
                doc.add_paragraph(f'Initial Notes: {notes}')

    def _add_programming_summary(self, doc: Document, df: pd.DataFrame, df_initial: pd.DataFrame, df_final: pd.DataFrame) -> None:
        """Add programming summary with session statistics."""
        from docx.shared import Pt

        doc.add_heading('Programming Summary', level=1)

        if df is None or df.empty:
            doc.add_paragraph("No session data available.")
            return

        # Session duration (from first to last timestamp)
        duration_str = "N/A"
        try:
            if 'time' in df.columns and 'date' in df.columns:
                timestamps = pd.to_datetime(
                    df['date'].astype(str) + ' ' + df['time'].astype(str),
                    errors='coerce',
                ).dropna()
            elif 'time' in df.columns:
                timestamps = pd.to_datetime(df['time'], errors='coerce').dropna()
            else:
                timestamps = pd.Series(dtype='datetime64[ns]')
            if len(timestamps) >= 2:
                duration = timestamps.max() - timestamps.min()
                total_mins = int(duration.total_seconds() / 60)
                if total_mins >= 60:
                    hours = total_mins // 60
                    mins = total_mins % 60
                    duration_str = f"{hours}h {mins}min"
                else:
                    duration_str = f"{total_mins} min"
        except Exception:
            pass

        # Number of configurations tested
        df_normalized = self._normalize_block_id_column(df)
        num_configs = 0
        if 'block_id' in df_normalized.columns:
            num_configs = df_normalized['block_id'].nunique()

        # Parameter ranges per side (Left / Right)
        def _param_range(series):
            vals = pd.to_numeric(series, errors='coerce').dropna()
            if len(vals) == 0:
                return "N/A"
            if vals.min() == vals.max():
                return f"{vals.min()}"
            return vals.min(), vals.max()

        amp_l = amp_r = freq_l = freq_r = pw_l = pw_r = "N/A"
        try:
            for prefix, side_label in [('left_', 'L'), ('right_', 'R')]:
                amp_col = f"{prefix}amplitude"
                freq_col = f"{prefix}stim_freq"
                pw_col = f"{prefix}pulse_width"
                if amp_col in df.columns:
                    r = _param_range(df[amp_col])
                    val = f"{r[0]:.1f} - {r[1]:.1f} mA" if isinstance(r, tuple) else (f"{float(r):.1f} mA" if r != "N/A" else r)
                    if side_label == 'L':
                        amp_l = val
                    else:
                        amp_r = val
                if freq_col in df.columns:
                    r = _param_range(df[freq_col])
                    val = f"{r[0]:.0f} - {r[1]:.0f} Hz" if isinstance(r, tuple) else (f"{float(r):.0f} Hz" if r != "N/A" else r)
                    if side_label == 'L':
                        freq_l = val
                    else:
                        freq_r = val
                if pw_col in df.columns:
                    r = _param_range(df[pw_col])
                    val = f"{r[0]:.0f} - {r[1]:.0f} µs" if isinstance(r, tuple) else (f"{float(r):.0f} µs" if r != "N/A" else r)
                    if side_label == 'L':
                        pw_l = val
                    else:
                        pw_r = val
        except Exception:
            pass

        # Add summary paragraphs
        summary_items = [
            f"Session Duration: {duration_str}",
            f"Configurations Tested: {num_configs}",
            f"Amplitude Range:  L: {amp_l}  |  R: {amp_r}",
            f"Frequency Range:  L: {freq_l}  |  R: {freq_r}",
            f"Pulse Width Range:  L: {pw_l}  |  R: {pw_r}",
        ]

        for item in summary_items:
            para = doc.add_paragraph(item)
            for run in para.runs:
                run.font.size = Pt(11)

        doc.add_paragraph("")

    def _create_lateral_table_data(self, df):
        """
        Create lateral table structure for Word and PDF exports.

        Returns DataFrame with lateral structure:
        - Left side parameters in first row
        - Right side parameters in second row
        - Non-lateral data merged vertically
        - Multiple scales from same block grouped in single cell
        """
        if df.empty:
            return df

        df = self._normalize_block_id_column(df)

        # Group by block_id to consolidate multiple scales
        if 'block_id' in df.columns:
            grouped = df.groupby('block_id', sort=False, dropna=False)
        else:
            grouped = [(0, df)]

        # Create new lateral structure
        lateral_data = []

        # Process each block
        for block_id, block_df in grouped:
            # Get first row to extract common values
            first_row = block_df.iloc[0]

            # Collect all scales for this block
            scale_pairs = []
            seen_pairs = set()
            for _, row in block_df.iterrows():
                sname = row.get('scale_name', '')
                sval = row.get('scale_value', '')
                if pd.notna(sname) and str(sname).strip():
                    name_str = str(sname).strip()
                    val_str = str(sval) if pd.notna(sval) else ''
                    key = (name_str, val_str)
                    if key not in seen_pairs:
                        seen_pairs.add(key)
                        scale_pairs.append(key)

            scale_names = [p[0] for p in scale_pairs]
            scale_values = [p[1] for p in scale_pairs]

            # Join multiple scales with newlines for internal separation
            combined_scale_name = '\n'.join(scale_names) if scale_names else ''
            combined_scale_value = '\n'.join(scale_values) if scale_values else ''

            # Left side row
            left_row = {}
            right_row = {}

            # Keep block_id in the output for styling logic (excluded from display columns later)
            left_row['block_id'] = block_id
            right_row['block_id'] = block_id

            # Common columns (non-lateral) - use combined scales with internal lines
            left_row['group_ID'] = first_row.get('group_ID', '')
            left_row['scale_name'] = combined_scale_name
            left_row['scale_value'] = combined_scale_value
            left_row['notes'] = first_row.get('notes', '')

            right_row['group_ID'] = first_row.get('group_ID', '')
            right_row['scale_name'] = combined_scale_name
            right_row['scale_value'] = combined_scale_value
            right_row['notes'] = first_row.get('notes', '')

            # Lateral columns - map to generic names
            lateral_mappings = {
                'left_stim_freq': 'frequency',
                'left_cathode': 'cathode',
                'left_anode': 'anode',
                'left_amplitude': 'amplitude',
                'left_pulse_width': 'pulse_width',
                'right_stim_freq': 'frequency',
                'right_cathode': 'cathode',
                'right_anode': 'anode',
                'right_amplitude': 'amplitude',
                'right_pulse_width': 'pulse_width',
            }

            # Left side parameters
            for left_col, generic_col in lateral_mappings.items():
                if left_col.startswith('left_'):
                    left_row[generic_col] = first_row.get(left_col, '')

            # Right side parameters
            for right_col, generic_col in lateral_mappings.items():
                if right_col.startswith('right_'):
                    right_row[generic_col] = first_row.get(right_col, '')

            # Add lateral indicator
            left_row['laterality'] = 'L'
            right_row['laterality'] = 'R'

            lateral_data.append(left_row)
            lateral_data.append(right_row)

        return pd.DataFrame(lateral_data)

    def _add_session_data_table(self, doc: Document, df_table: pd.DataFrame) -> None:
        """Add the lateral session-data table to the Word document."""
        doc.add_heading('Session Data', level=1)

        if df_table is None or df_table.empty:
            return

        df_table = self._normalize_block_id_column(df_table)
        lateral_df = self._create_lateral_table_data(df_table)

        # Chart BEFORE the table
        self._add_scales_timeline_chart(doc, lateral_df)

        columns_to_exclude = ['date', 'time', 'onset', 'block_id', 'session_ID', 'is_initial', 'electrode_model']
        display_columns = [col for col in lateral_df.columns if col not in columns_to_exclude]

        lateral_cols = ['laterality', 'frequency', 'anode', 'cathode', 'amplitude', 'pulse_width']
        common_cols = ['group_ID', 'scale_name', 'scale_value', 'notes']

        lateral_cols = [col for col in lateral_cols if col in display_columns]
        common_cols = [col for col in common_cols if col in display_columns]
        ordered_columns = lateral_cols + common_cols

        table = doc.add_table(rows=lateral_df.shape[0] + 1, cols=len(ordered_columns))
        table.style = 'Table Grid'
        table.autofit = False

        # Define column widths in inches
        section = doc.sections[0]
        page_width_inches = (section.page_width - section.left_margin - section.right_margin) / 914400

        base_in = {
            'laterality': 0.30,
            'group_ID': 0.40,
            'frequency': 0.50,
            'anode': 0.45,
            'cathode': 0.60,
            'amplitude': 0.60,
            'pulse_width': 0.50,
            'scale_name': 1.10,
            'scale_value': 0.60,
        }
        widths_in = [base_in.get(c, 0.5) for c in ordered_columns]
        if 'notes' in ordered_columns:
            notes_idx = ordered_columns.index('notes')
            used = sum(w for j, w in enumerate(widths_in) if j != notes_idx)
            widths_in[notes_idx] = max(2.5, page_width_inches - used)

        # Apply widths to each cell in every row (required for python-docx)
        widths_twips = [Inches(max(0.25, w)) for w in widths_in]
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = widths_twips[idx]

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(ordered_columns):
            hdr_cells[i].text = self._column_header(col_name)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Find best and second-best blocks for green highlighting
        best_block_ids, second_best_ids = self._find_best_and_second_best_blocks(lateral_df)

        prev_block_id = None
        scale_name_idx = ordered_columns.index('scale_name') if 'scale_name' in ordered_columns else -1
        scale_value_idx = ordered_columns.index('scale_value') if 'scale_value' in ordered_columns else -1

        for i, (_, row) in enumerate(lateral_df.iterrows()):
            row_cells = table.rows[i + 1].cells

            # Highlight best block(s) with darker green, second-best with lighter green
            current_block_id = row.get('block_id', None)
            if best_block_ids and current_block_id in best_block_ids:
                self._highlight_cells_green(row_cells, intensity="best")
            elif second_best_ids and current_block_id in second_best_ids:
                self._highlight_cells_green(row_cells, intensity="second")

            current_block_id = row.get('block_id', None)
            if (
                prev_block_id is not None
                and current_block_id != prev_block_id
                and row.get('laterality') == 'L'
            ):
                for cell in row_cells:
                    self._set_cell_border_top(cell, sz=24)
            prev_block_id = current_block_id

            scale_name_lines = None
            scale_value_lines = None
            if scale_name_idx >= 0 and scale_value_idx >= 0:
                try:
                    raw_sn = row.get('scale_name', '')
                    raw_sv = row.get('scale_value', '')
                    sn_text = str(raw_sn) if pd.notna(raw_sn) else ''
                    sv_text = str(raw_sv) if pd.notna(raw_sv) else ''

                    # Filter out NaN scales from display - remove both name and value
                    if sv_text == 'NaN' or 'NaN' in sv_text:
                        sn_parts = [s for s in sn_text.split('\n') if s.strip()]
                        sv_parts = [s for s in sv_text.split('\n') if s.strip()]
                        # Remove corresponding name-value pairs where value is NaN
                        filtered_names = []
                        filtered_values = []
                        for name, val in zip(sn_parts, sv_parts, strict=False):
                            if val != 'NaN' and val.strip() != 'NaN':
                                filtered_names.append(name)
                                filtered_values.append(val)
                        sn_text = '\n'.join(filtered_names)
                        sv_text = '\n'.join(filtered_values)

                    scale_name_lines = sn_text.split('\n') if sn_text else ['']
                    scale_value_lines = sv_text.split('\n') if sv_text else ['']
                    max_len = max(len(scale_name_lines), len(scale_value_lines))
                    scale_name_lines += [''] * (max_len - len(scale_name_lines))
                    scale_value_lines += [''] * (max_len - len(scale_value_lines))
                except Exception:
                    scale_name_lines = None
                    scale_value_lines = None

            for j, col in enumerate(ordered_columns):
                if col not in row:
                    continue

                # Format numeric values: use int for frequency/pulse_width when no decimals
                cell_value = str(row[col]) if pd.notna(row[col]) else ''
                if col in ['frequency', 'pulse_width']:
                    try:
                        val = float(row[col])
                        if val.is_integer():
                            cell_value = str(int(val))
                    except (ValueError, TypeError):
                        pass

                if col in common_cols:
                    merged_target_cell = None
                    did_merge = False
                    if row.get('laterality') == 'R' and i > 0:
                        prev_cell = table.rows[i].cells[j]
                        prev_cell.merge(row_cells[j])
                        row_cells[j].text = ''
                        merged_target_cell = prev_cell
                        did_merge = True

                    target_cell = merged_target_cell if did_merge and merged_target_cell is not None else row_cells[j]
                    if col == 'scale_name' and scale_name_lines is not None and len(scale_name_lines) > 1:
                        target_cell.text = "\n".join(scale_name_lines)
                    elif col == 'scale_value' and scale_value_lines is not None and len(scale_value_lines) > 1:
                        target_cell.text = "\n".join(scale_value_lines)
                    else:
                        target_cell.text = cell_value
                elif col == 'cathode' and '_' in cell_value:
                    # Multi-contact cathode: show stacked with Total label
                    contacts = cell_value.replace('_', '\n')
                    row_cells[j].text = contacts + '\nTotal'
                elif col == 'amplitude' and '_' in cell_value:
                    # Multi-contact amplitude: show stacked values with total
                    parts = cell_value.split('_')
                    try:
                        # Validate all parts are numbers and calculate total
                        values = [float(p) for p in parts]
                        total = sum(values)
                        total_str = f"{total:.2f}".rstrip('0').rstrip('.')
                        row_cells[j].text = '\n'.join(parts) + f'\n{total_str}'
                    except (ValueError, TypeError):
                        row_cells[j].text = cell_value
                else:
                    row_cells[j].text = cell_value

        # Add legend and clinical disclaimer below table
        self._add_table_legend(doc, best_block_ids, second_best_ids)

    def _add_table_legend(self, doc: Document, best_ids: list, second_ids: list) -> None:
        """Add color legend and clinical disclaimer below the session data table."""
        from docx.shared import Pt, RGBColor

        # Only add legend if there are highlighted blocks
        if not best_ids and not second_ids:
            return

        doc.add_paragraph()  # spacing

        # Legend paragraph
        legend_para = doc.add_paragraph()
        legend_para.add_run("Legend: ").bold = True

        if best_ids:
            best_run = legend_para.add_run("■ ")
            best_run.font.color.rgb = RGBColor(0xC3, 0xE6, 0xCB)
            legend_para.add_run("Optimal configuration    ")

        if second_ids:
            second_run = legend_para.add_run("■ ")
            second_run.font.color.rgb = RGBColor(0xE8, 0xF5, 0xE9)
            legend_para.add_run("Second-best configuration")

        # Show target values used for optimization
        if self.scale_optimization_prefs:
            targets_para = doc.add_paragraph()
            targets_para.add_run("Scale targets: ").bold = True
            target_parts = []
            for pref in self.scale_optimization_prefs:
                if len(pref) >= 5:
                    name, smin, smax, mode, custom_val = pref
                    if mode == "ignore":
                        continue
                    elif mode == "min":
                        target_parts.append(f"{name}: min")
                    elif mode == "max":
                        target_parts.append(f"{name}: max")
                    elif mode == "custom":
                        target_parts.append(f"{name}: {custom_val}")
            if target_parts:
                targets_para.add_run("; ".join(target_parts))
                for run in targets_para.runs:
                    run.font.size = Pt(9)

        # Clinical disclaimer
        disclaimer_para = doc.add_paragraph()
        disclaimer_run = disclaimer_para.add_run(
            "Note: The highlighted rows are derived exclusively from the recorded "
            "session scale values and represent a computational ranking intended "
            "solely as a reference. This color-coded indication does not constitute "
            "clinical guidance."
        )
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.font.italic = True

    def _add_scales_timeline_chart(self, doc: Document, lateral_df: pd.DataFrame) -> None:
        """Add a rainbow-colored timeline chart of session scales with a general index line."""
        import math as _math
        from io import BytesIO

        # Guard: need valid input
        if lateral_df is None or lateral_df.empty:
            doc.add_paragraph('No session data available for chart.')
            return
        if 'scale_name' not in lateral_df.columns or 'scale_value' not in lateral_df.columns:
            doc.add_paragraph('No scale columns found in session data.')
            return
        if 'block_id' not in lateral_df.columns:
            doc.add_paragraph('No block ID column found in session data.')
            return

        # Use L rows only to avoid duplicates
        if 'laterality' in lateral_df.columns:
            df_l = lateral_df[lateral_df['laterality'] == 'L'].copy()
        else:
            df_l = lateral_df.copy()
        if df_l.empty:
            df_l = lateral_df.drop_duplicates(subset=['block_id']).copy()

        # Collect scale values per block
        scale_data = {}  # scale_name -> {block_id: value}
        for _, row in df_l.iterrows():
            try:
                block_id = int(row.get('block_id', 0))
            except (ValueError, TypeError):
                continue
            names = str(row.get('scale_name', '') or '').split('\n')
            values = str(row.get('scale_value', '') or '').split('\n')
            for i, name in enumerate(names):
                name = name.strip()
                if not name:
                    continue
                val_str = values[i].strip() if i < len(values) else ''
                try:
                    val = float(val_str)
                except ValueError:
                    continue
                if _math.isnan(val):
                    continue
                scale_data.setdefault(name, {})[block_id] = val

        if not scale_data:
            doc.add_paragraph('No numeric scale values recorded for this session.')
            return

        all_blocks = sorted({b for pts in scale_data.values() for b in pts})
        if not all_blocks:
            doc.add_paragraph('No configuration blocks with scale data found.')
            return

        try:
            import pyqtgraph as pg
            from PySide6.QtCore import QBuffer, QIODevice, Qt
            from PySide6.QtGui import QBrush, QColor, QFont, QPen

            pg.setConfigOptions(useOpenGL=False, antialias=True)

            n_scales = len(scale_data)
            rainbow = [QColor.fromHsvF(i / max(n_scales, 1), 0.85, 0.85)
                        for i in range(n_scales)]

            has_index = n_scales >= 2
            win = pg.GraphicsLayoutWidget()
            win.setBackground('w')
            win.resize(1050, 500)  # Single plot, larger for right-side legend

            # --- Main scales chart with General Index on same plot ---
            p1 = win.addPlot(row=0, col=0)
            p1.setTitle('Session Scales Timeline', color='k', size='14pt')
            p1.setLabel('left', 'Scale Value', color='k', size='14pt', font='Arial')
            p1.setLabel('bottom', 'Block', color='k', size='14pt', font='Arial')
            p1.getAxis('left').setStyle(tickFont=QFont('Arial', 10))
            p1.getAxis('bottom').setStyle(tickFont=QFont('Arial', 10))
            p1.showGrid(x=True, y=True, alpha=0.3)
            # Legend on right side external - increase offset and add background
            legend = p1.addLegend(offset=(1.15, 0.5), pen=QPen(Qt.black, 1), brush=QBrush(Qt.white))
            legend.setLabelTextColor('k')

            # Plot individual scales with original values (no normalization)
            for idx, (sname, pts) in enumerate(scale_data.items()):
                c = rainbow[idx]
                xs = sorted(pts.keys())
                ys = [pts[x] for x in xs]
                p1.plot(xs, ys,
                        pen=pg.mkPen(c, width=2),
                        symbol='o', symbolPen=pg.mkPen(c, width=1),
                        symbolBrush=pg.mkBrush(c), symbolSize=8,
                        name=sname)

            # --- General Index on same plot (if >= 2 scales) ---
            if has_index:
                # Create scale targets dictionary from preferences
                scale_targets = {}
                if self.scale_optimization_prefs:
                    for pref in self.scale_optimization_prefs:
                        if len(pref) >= 5:
                            name, smin, smax, mode, custom_val = pref
                            if mode == "min":
                                scale_targets[name] = {"type": "min", "value": smin}
                            elif mode == "max":
                                scale_targets[name] = {"type": "max", "value": smax}
                            elif mode == "custom":
                                try:
                                    scale_targets[name] = {"type": "custom", "value": float(custom_val)}
                                except ValueError:
                                    scale_targets[name] = {"type": "custom", "value": 0.0}

                index_vals = {}
                for b in all_blocks:
                    weighted_scores = []
                    weights = []

                    for scale_name in scale_data:
                        if b in scale_data[scale_name]:
                            original_value = scale_data[scale_name][b]

                            # Get target for this scale
                            if scale_name in scale_targets:
                                target_info = scale_targets[scale_name]
                                target_type = target_info["type"]
                                target_value = target_info["value"]

                                # Calculate distance from target (lower is better)
                                if target_type == "min":
                                    # For minimization: lower values are better
                                    distance = original_value
                                    max_possible = max(scale_data[scale_name].values())
                                    # Normalize: 0 = at target (min), 1 = worst (max)
                                    normalized_score = distance / max_possible if max_possible > 0 else 0
                                elif target_type == "max":
                                    # For maximization: higher values are better
                                    distance = max(scale_data[scale_name].values()) - original_value
                                    max_possible = max(scale_data[scale_name].values()) - min(scale_data[scale_name].values())
                                    # Normalize: 0 = at target (max), 1 = worst (min)
                                    normalized_score = distance / max_possible if max_possible > 0 else 0
                                elif target_type == "custom":
                                    # For custom target: absolute distance from target
                                    distance = abs(original_value - target_value)
                                    max_distance = max(abs(v - target_value) for v in scale_data[scale_name].values())
                                    # Normalize: 0 = at target, 1 = worst
                                    normalized_score = distance / max_distance if max_distance > 0 else 0

                                # Convert to proximity score (higher is better)
                                proximity_score = 1.0 - normalized_score
                                weighted_scores.append(proximity_score)
                                weights.append(1.0)  # Equal weight for now
                            else:
                                # No target defined: use neutral score
                                weighted_scores.append(0.5)  # Neutral middle value
                                weights.append(0.5)  # Lower weight for scales without targets

                    if weighted_scores and weights:
                        # Calculate weighted average of proximity scores
                        total_weight = sum(weights)
                        if total_weight > 0:
                            index_vals[b] = sum(w * s for w, s in zip(weights, weighted_scores, strict=False)) / total_weight
                        else:
                            index_vals[b] = 0.5  # Default neutral value

                if index_vals:
                    ix = sorted(index_vals.keys())
                    iy = [index_vals[x] for x in ix]
                    # Thicker black line for General Index
                    p1.plot(ix, iy,
                            pen=pg.mkPen('k', width=5),
                            symbol='d', symbolPen='k', symbolBrush='k',
                            symbolSize=10, name='General Index')

            # --- Export to PNG → Word ---
            pixmap = win.grab()
            qbuf = QBuffer()
            qbuf.open(QIODevice.OpenModeFlag.WriteOnly)
            pixmap.save(qbuf, 'PNG')
            qbuf.close()
            img_buf = BytesIO(bytes(qbuf.data()))
            doc.add_picture(img_buf, width=Inches(6))
            doc.add_paragraph()
            img_buf.close()
            win.close()
            del win

        except Exception as e:
            doc.add_paragraph(f'Chart generation error: {e}')

    def _column_header(self, col: str) -> str:
        """Map an internal column name to a human-readable table header."""
        placeholder_map = {
            "scale_name": PLACEHOLDERS.get("scale_name"),
            "scale_value": PLACEHOLDERS.get("scale_value"),
            "frequency": PLACEHOLDERS.get("frequency"),
            "anode": "+",
            "cathode": "-",
            "amplitude": PLACEHOLDERS.get("amplitude"),
            "pulse_width": PLACEHOLDERS.get("pulse_width"),
            "group_ID": "grp",
            "laterality": "",
        }
        if col in placeholder_map and placeholder_map[col] is not None:
            return str(placeholder_map[col])
        return str(col).replace('_', ' ').title()

    def _pick_latest_row(self, df: pd.DataFrame) -> pd.Series | None:
        """Return the row with the highest block_id, or the last row."""
        if df is None or df.empty:
            return None
        if "block_id" in df.columns:
            try:
                bid = pd.to_numeric(df["block_id"], errors="coerce")
                return df.loc[bid.idxmax()]
            except Exception:
                return df.iloc[-1]
        return df.iloc[-1]

    def _find_best_and_second_best_blocks(self, lateral_df: pd.DataFrame) -> tuple:
        """
        Find block_ids with the best and second-best scores based on optimization preferences.

        Returns a tuple: (best_block_ids, second_best_block_ids)
        Each is a list (may have multiple if tied).
        """
        if lateral_df is None or lateral_df.empty:
            return [], []
        if 'block_id' not in lateral_df.columns or 'scale_value' not in lateral_df.columns:
            return [], []
        if 'scale_name' not in lateral_df.columns:
            return [], []

        try:
            # Build preference lookup: scale_name -> (mode, custom_value)
            pref_lookup = {}
            for pref in self.scale_optimization_prefs:
                if len(pref) >= 5:
                    name, _, _, mode, custom_val = pref
                    pref_lookup[name.strip().lower()] = (mode, custom_val)

            # Get unique blocks (use only L rows to avoid double counting)
            df_l = lateral_df[lateral_df.get('laterality', '') == 'L'].copy()
            if df_l.empty:
                df_l = lateral_df.drop_duplicates(subset=['block_id']).copy()

            block_scores = {}
            for _, row in df_l.iterrows():
                block_id = row.get('block_id')
                if block_id is None:
                    continue

                scale_name_str = str(row.get('scale_name', '') or '')
                scale_val_str = str(row.get('scale_value', '') or '')

                names = scale_name_str.split('\n')
                values = scale_val_str.split('\n')

                # Calculate weighted score for this block
                total_score = 0.0
                has_value = False

                for i, val_line in enumerate(values):
                    val_line = val_line.strip()
                    if not val_line:
                        continue

                    try:
                        val = float(val_line)
                    except ValueError:
                        continue

                    import math as _math
                    if _math.isnan(val):
                        continue

                    # Get the corresponding scale name
                    scale_name = names[i].strip().lower() if i < len(names) else ""
                    mode, custom_val = pref_lookup.get(scale_name, ("min", ""))

                    if mode == "ignore":
                        continue  # Skip this scale

                    has_value = True

                    if mode in ("low", "min"):
                        # Lower is better - use value directly as score (lower = better)
                        total_score += val
                    elif mode in ("high", "max"):
                        # Higher is better - negate so lower score = better
                        total_score -= val
                    elif mode == "custom":
                        # Closer to custom value is better - use absolute distance
                        try:
                            target = float(custom_val) if custom_val else 0.0
                            total_score += abs(val - target)
                        except ValueError:
                            total_score += val

                if has_value:
                    block_scores[block_id] = total_score

            if not block_scores:
                return [], []

            # Sort unique scores
            unique_scores = sorted(set(block_scores.values()))

            # Best blocks (lowest score)
            best_score = unique_scores[0]
            best_blocks = [bid for bid, score in block_scores.items() if score == best_score]

            # Second best blocks (second lowest score, if exists)
            second_best_blocks = []
            if len(unique_scores) > 1:
                second_score = unique_scores[1]
                second_best_blocks = [bid for bid, score in block_scores.items() if score == second_score]

            return best_blocks, second_best_blocks

        except Exception:
            return [], []

    def _highlight_cells_green(self, row_cells, intensity: str = "best") -> None:
        """
        Apply green background to all cells in a row.

        Args:
            row_cells: List of cells to highlight
            intensity: "best" for darker green, "second" for lighter green
        """
        # Best = darker green, Second = lighter green
        color = 'C3E6CB' if intensity == "best" else 'E8F5E9'
        for cell in row_cells:
            try:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), color)
                cell._tc.get_or_add_tcPr().append(shading_elm)
            except Exception:
                pass

    def _set_cell_border_top(self, cell, sz=12):
        """Set top border of a cell to specified size (in eighths of a point)."""
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr() # noqa: N806
            tcBorders = OxmlElement('w:tcBorders') # noqa: N806
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), str(sz))
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '000000')
            tcBorders.append(top)
            tcPr.append(tcBorders)
        except Exception:
            pass

    def _set_paragraph_bottom_border(self, paragraph, sz: int = 6, color: str = '000000') -> None:
        """Draw a bottom border line under a Word paragraph."""
        try:
            pPr = paragraph._p.get_or_add_pPr() # noqa: N806
            pBdr = OxmlElement('w:pBdr') # noqa: N806
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(sz))
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), str(color))
            pBdr.append(bottom)
            pPr.append(pBdr)
        except Exception:
            pass

    def _write_multiline_cell_with_dividers(self, cell, lines: list, divider_sz: int = 12, divider_color: str = '000000') -> None:
        """Write each line as its own paragraph and draw a full-width divider under each line except the last."""
        try:
            cell.text = ''
            if not lines:
                return

            p0 = cell.paragraphs[0]
            p0.text = str(lines[0])
            if len(lines) > 1:
                self._set_paragraph_bottom_border(p0, sz=divider_sz, color=divider_color)

            for k in range(1, len(lines)):
                p = cell.add_paragraph(str(lines[k]))
                if k < len(lines) - 1:
                    self._set_paragraph_bottom_border(p, sz=divider_sz, color=divider_color)
        except Exception:
            try:
                cell.text = "\n".join([str(x) for x in (lines or [])])
            except Exception:
                pass

    def _apply_contact_tokens_to_canvas(self, canvas: ElectrodeCanvas, anode_text: str, cathode_text: str) -> None:
        """Parse anode/cathode token strings and set the corresponding canvas states."""
        model = canvas.model
        if not model:
            return

        canvas.contact_states.clear()
        canvas.case_state = ContactState.OFF

        def apply_tokens(text: str, state: ContactState) -> None:
            if not text:
                return
            for token in str(text).split("_"):
                token = token.strip()
                if not token:
                    continue
                if token == "case":
                    canvas.case_state = state
                    continue

                if token.startswith("E") and len(token) >= 2:
                    try:
                        if token[-1].isalpha():
                            idx = int(token[1:-1])
                            seg_char = token[-1].lower()
                            seg_map = {"a": 0, "b": 1, "c": 2}
                            if seg_char in seg_map:
                                canvas.contact_states[(idx, seg_map[seg_char])] = state
                        else:
                            idx = int(token[1:])
                            if model.is_directional:
                                for seg in range(3):
                                    canvas.contact_states[(idx, seg)] = state
                            else:
                                canvas.contact_states[(idx, 0)] = state
                    except Exception:
                        continue

        apply_tokens(anode_text, ContactState.ANODIC)
        apply_tokens(cathode_text, ContactState.CATHODIC)
        canvas.update()

    def _render_electrode_png(
        self,
        model_name: str,
        anode_text: str,
        cathode_text: str,
        target_size_px: tuple[int, int] = (440, 900),
    ) -> str | None:
        model = ELECTRODE_MODELS.get(model_name)
        if not model:
            return None

        canvas = ElectrodeCanvas()
        canvas.set_model(model)
        canvas.resize(*target_size_px)
        try:
            canvas.set_export_mode(True)
        except Exception:
            pass
        self._apply_contact_tokens_to_canvas(canvas, anode_text, cathode_text)

        # Force white background by temporarily overriding paintEvent
        original_paint = canvas.paintEvent
        def white_bg_paint(event):
            painter = QPainter(canvas)
            painter.fillRect(canvas.rect(), Qt.white)
            original_paint(event)
        canvas.paintEvent = white_bg_paint

        pixmap = QPixmap(canvas.size())
        pixmap.fill(Qt.white)
        canvas.render(pixmap)

        # Crop white borders
        image = pixmap.toImage()
        from PySide6.QtGui import QColor as _QColor
        # Find bounding box of non-white content
        left, top, right, bottom = image.width(), image.height(), 0, 0
        white_rgb = _QColor(Qt.white).rgb()
        for y in range(image.height()):
            for x in range(image.width()):
                if image.pixel(x, y) != white_rgb:
                    left = min(left, x)
                    top = min(top, y)
                    right = max(right, x)
                    bottom = max(bottom, y)
        if right > left and bottom > top:
            margin = 20  # small margin in pixels
            left = max(0, left - margin)
            top = max(0, top - margin)
            right = min(image.width() - 1, right + margin)
            bottom = min(image.height() - 1, bottom + margin)
            cropped = pixmap.copy(left, top, right - left + 1, bottom - top + 1)
        else:
            cropped = pixmap

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp.close()
        cropped.save(tmp.name, "PNG")
        return tmp.name

    def _add_electrode_config_section(self, doc: Document, df: pd.DataFrame, df_initial: pd.DataFrame) -> None:
        """Add the initial vs final electrode configuration images to the document."""
        if df is None or df.empty:
            return
        if "session_ID" not in df.columns or "is_initial" not in df.columns:
            return

        dfc = df.copy()
        dfc["session_ID"] = pd.to_numeric(dfc["session_ID"], errors="coerce")
        dfc["is_initial"] = pd.to_numeric(dfc["is_initial"], errors="coerce").fillna(0).astype(int)

        df_init = dfc[dfc["is_initial"] == 1]
        df_final = dfc[dfc["is_initial"] == 0]

        if df_init.empty or df_final.empty:
            return

        init_session = int(dfc.loc[df_init.index, "session_ID"].max())
        final_session = int(dfc.loc[df_final.index, "session_ID"].max())

        init_row = self._pick_latest_row(df_init[df_init["session_ID"] == init_session])
        final_row = self._pick_latest_row(df_final[df_final["session_ID"] == final_session])

        if init_row is None or final_row is None:
            return

        init_model = str(init_row.get("electrode_model", "") or "")
        final_model = str(final_row.get("electrode_model", "") or "")
        if not init_model or not final_model:
            return

        paths = {
            "Init L": self._render_electrode_png(init_model, str(init_row.get("left_anode", "") or ""), str(init_row.get("left_cathode", "") or "")),
            "Init R": self._render_electrode_png(init_model, str(init_row.get("right_anode", "") or ""), str(init_row.get("right_cathode", "") or "")),
            "Final L": self._render_electrode_png(final_model, str(final_row.get("left_anode", "") or ""), str(final_row.get("left_cathode", "") or "")),
            "Final R": self._render_electrode_png(final_model, str(final_row.get("right_anode", "") or ""), str(final_row.get("right_cathode", "") or "")),
        }

        if not all(paths.values()):
            return

        doc.add_heading("Electrode Configurations", level=1)

        # Add electrode model info
        latest_init = self._pick_latest_session_row(df_initial)
        model = str(latest_init.get('electrode_model', '') or '')
        manufacturer = self._get_manufacturer_for_model(model)
        if model:
            if manufacturer:
                doc.add_paragraph(f'Electrode model: {manufacturer} | {model}')
            else:
                doc.add_paragraph(f'Electrode model: {model}')

        # 4 columns x 4 rows table, no borders
        # Row 0: "Initial Settings" (merged cols 0-1), "Final Settings" (merged cols 2-3)
        # Row 1: Left, Right, Left, Right
        # Row 2: Anode/Cathode config text
        # Row 3: PNG images
        t = doc.add_table(rows=4, cols=4)
        t.autofit = False

        # Remove all borders
        tbl = t._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr() # noqa: N806
        borders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        tblPr.append(borders)

        # Row 0: merged headers
        cell_init = t.cell(0, 0).merge(t.cell(0, 1))
        cell_init.text = "Initial Settings"
        for p in cell_init.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

        cell_final = t.cell(0, 2).merge(t.cell(0, 3))
        cell_final.text = "Final Settings"
        for p in cell_final.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

        # Row 1: Left / Right labels
        for c, label in enumerate(["Left", "Right", "Left", "Right"]):
            cell = t.cell(1, c)
            cell.text = label
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Captions in order: Init L, Init R, Final L, Final R
        all_captions = [
            (str(init_row.get("left_anode", "") or ""), str(init_row.get("left_cathode", "") or "")),
            (str(init_row.get("right_anode", "") or ""), str(init_row.get("right_cathode", "") or "")),
            (str(final_row.get("left_anode", "") or ""), str(final_row.get("left_cathode", "") or "")),
            (str(final_row.get("right_anode", "") or ""), str(final_row.get("right_cathode", "") or "")),
        ]
        all_img_paths = [paths["Init L"], paths["Init R"], paths["Final L"], paths["Final R"]]

        # Row 2: Config text
        for c, (anode_txt, cathode_txt) in enumerate(all_captions):
            cell = t.cell(2, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.text = f"Anode: {anode_txt}\nCathode: {cathode_txt}".strip()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Row 3: PNG images
        for c, img_path in enumerate(all_img_paths):
            cell = t.cell(3, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(1.15))

        for pth in paths.values():
            try:
                os.unlink(pth)
            except Exception:
                pass

    def export_to_pdf(self, parent: QWidget | None = None, sections=None) -> bool:
        """Export session data to PDF by generating a Word report and converting it."""
        try:
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent,
                    "No Session Data",
                    "No session file is currently open. Please start a session first.",
                )
                return False

            default_filename = self._generate_bids_report_filename('.pdf')
            start_dir = os.path.dirname(getattr(self.session_data, 'file_path', '') or '')
            start_path = os.path.join(start_dir, default_filename) if start_dir else default_filename

            from PySide6.QtWidgets import QFileDialog

            pdf_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Session Report",
                start_path,
                "PDF Files (*.pdf);;All Files (*)",
            )

            if not pdf_path:
                return False

            if not pdf_path.endswith(".pdf"):
                pdf_path += ".pdf"

            docx_path = os.path.splitext(pdf_path)[0] + "_tmp.docx"

            ok = self._export_to_word_path(docx_path, sections=sections)
            if not ok:
                return False

            try:
                self._convert_docx_to_pdf(docx_path, pdf_path)
            finally:
                try:
                    os.unlink(docx_path)
                except Exception:
                    pass

            self._show_transient_message(
                parent,
                "Export Completed",
                f"Report saved successfully:\n{pdf_path}",
                msecs=2000,
            )
            return True

        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export session data to PDF:\n{str(e)}",
            )
            return False

    def export_to_word(self, parent: QWidget | None = None, sections=None) -> bool:
        """
        Export session data to Word format.

        Args:
            parent: Parent widget for dialog display
            sections: List of section keys to include

        Returns:
            True if export was successful, False otherwise
        """
        try:
            # Get the current session data
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent,
                    "No Session Data",
                    "No session file is currently open. Please start a session first."
                )
                return False

            # Generate BIDS-friendly default filename from TSV path
            default_filename = self._generate_bids_report_filename('.docx')

            # Use same directory as TSV file
            start_dir = os.path.dirname(getattr(self.session_data, 'file_path', '') or '')
            start_path = os.path.join(start_dir, default_filename) if start_dir else default_filename

            # Get save location
            from PySide6.QtWidgets import QFileDialog
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Session Report",
                start_path,
                "Word Files (*.docx);;All Files (*)"
            )

            if not file_path:
                return False  # User cancelled

            # Ensure .docx extension
            if not file_path.endswith('.docx'):
                file_path += '.docx'

            ok = self._export_to_word_path(file_path, sections=sections)
            if not ok:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No session data has been recorded yet.",
                )
                return False

            self._show_transient_message(
                parent,
                "Export Completed",
                f"Report saved successfully:\n{file_path}",
                msecs=2000,
            )
            return True

        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export session data to Word:\n{str(e)}"
            )
            return False

    def _export_to_word_path(self, file_path: str, sections=None) -> bool:
        """Generate the Word report at an explicit path (used also by PDF export)."""
        df = self._read_session_data()
        if df is None or df.empty:
            return False

        df = df.copy()
        df = self._normalize_block_id_column(df)
        if "is_initial" in df.columns:
            df["is_initial"] = pd.to_numeric(df["is_initial"], errors="coerce").fillna(0).astype(int)

        df_initial = df[df.get("is_initial", 0) == 1] if "is_initial" in df.columns else df.iloc[0:0]
        df_table = df[df.get("is_initial", 0) != 1] if "is_initial" in df.columns else df

        doc = Document()

        section = doc.sections[0]
        section.left_margin = Inches(0.5)   # default  ~1.0
        section.right_margin = Inches(0.5)  # default  ~1.0
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        title = doc.add_heading('Clinical DBS Session Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Generated on: {datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S")} by {__app_name__} v{__version__}')

        patient_id, session_num = self._extract_bids_info_from_path()
        if patient_id or session_num:
            info_parts = []
            if patient_id:
                info_parts.append(f'Patient ID: {patient_id}')
            if session_num:
                info_parts.append(f'Session: {session_num}')
            doc.add_paragraph('    '.join(info_parts))

        # Determine which sections to include (default: all)
        all_keys = ["initial_notes", "session_data", "electrode_config", "programming_summary"]
        active = set(sections) if sections is not None else set(all_keys)

        if "initial_notes" in active:
            doc.add_paragraph('')
            self._add_summary_section(doc, df, df_initial, df_table)

        if "session_data" in active:
            doc.add_paragraph('')
            self._add_session_data_table(doc, df_table)

        if "electrode_config" in active:
            doc.add_paragraph('')
            self._add_electrode_config_section(doc, df, df_initial)

        if "programming_summary" in active:
            doc.add_paragraph('')
            self._add_programming_summary(doc, df, df_initial, df_table)

        doc.save(file_path)
        return True

    def _add_report_footer(self, doc: Document) -> None:
        """Add footer with patient ID and session number."""
        from docx.shared import Pt

        patient_id, session_num = self._extract_bids_info_from_path()

        if patient_id or session_num:
            doc.add_paragraph('')
            doc.add_paragraph('')

            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run("─" * 50)
            footer_run.font.size = Pt(8)

            info_para = doc.add_paragraph()
            if patient_id:
                info_para.add_run(f"Patient ID: {patient_id}    ")
            if session_num:
                info_para.add_run(f"Session: {session_num}")
            for run in info_para.runs:
                run.font.size = Pt(10)

    def _read_simple_annotations(self) -> list[tuple[str, str]]:
        """Read (timestamp, annotation) pairs from the simple annotations TSV."""
        file_path = getattr(self.session_data, "file_path", None)
        if not file_path and getattr(self.session_data, "tsv_file", None) is not None:
            try:
                file_path = self.session_data.tsv_file.name
            except Exception:
                file_path = None
        if not file_path or not os.path.exists(file_path):
            return []

        items: list[tuple[str, str]] = []
        try:
            with open(file_path, newline="", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f, delimiter="\t")
                for row in reader:
                    if not row:
                        continue

                    norm = {}
                    for k, v in row.items():
                        try:
                            kk = str(k).strip().lstrip("\ufeff").lower()
                        except Exception:
                            continue
                        norm[kk] = v

                    # Extract date and time separately
                    date_str = norm.get("date", "")
                    time_str = norm.get("time", "")

                    # Build timestamp: if both exist, combine; otherwise fallback to time/timestamp/date
                    if date_str and time_str:
                        t = f"{date_str} {time_str}"
                    else:
                        t = str(
                            norm.get("time", "")
                            or norm.get("timestamp", "")
                            or norm.get("date", "")
                            or ""
                        )

                    a = str(
                        norm.get("annotation", "")
                        or norm.get("note", "")
                        or norm.get("notes", "")
                        or norm.get("text", "")
                        or ""
                    )

                    # Fallback: if headers are unexpected, use the first/second column values
                    if not (t.strip() or a.strip()):
                        try:
                            vals = list(norm.values())
                            t = str(vals[0] or "") if len(vals) >= 1 else ""
                            a = str(vals[1] or "") if len(vals) >= 2 else ""
                        except Exception:
                            t, a = "", ""

                    if a.strip() or t.strip():
                        items.append((t, a))
        except Exception:
            return []
        return items

    def _export_annotations_to_word_path(self, file_path: str) -> bool:
        """Generate the annotations Word report at an explicit path."""
        annotations = self._read_simple_annotations()
        if not annotations:
            return False

        doc = Document()

        title = doc.add_heading("Annotations Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated on: {datetime.now().astimezone().strftime('%Y-%m-%d %H:%M:%S')} by {__app_name__} v{__version__}")
        doc.add_paragraph("")

        for t, a in annotations:
            line = f"{t}: {a}" if t.strip() else a
            doc.add_paragraph(line)

        doc.save(file_path)
        return True

    def export_annotations_to_word(self, parent: QWidget | None = None) -> bool:
        """Export simple annotations to a Word document."""
        try:
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent,
                    "No Session Data",
                    "No annotation file is currently open. Please open or create one first.",
                )
                return False

            if getattr(self.session_data, "tsv_file", None) is not None:
                try:
                    self.session_data.tsv_file.flush()
                except Exception:
                    pass

            from PySide6.QtWidgets import QFileDialog
            default_filename = self._generate_bids_report_filename('.docx')
            start_dir = os.path.dirname(getattr(self.session_data, "file_path", "") or "")
            start_path = os.path.join(start_dir, default_filename) if start_dir else default_filename
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Annotations Report",
                start_path,
                "Word Files (*.docx);;All Files (*)",
            )
            if not file_path:
                return False
            if not file_path.endswith(".docx"):
                file_path += ".docx"

            ok = self._export_annotations_to_word_path(file_path)
            if not ok:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No annotations have been recorded yet.",
                )
                return False

            self._show_transient_message(
                parent,
                "Export Completed",
                f"Report saved successfully:\n{file_path}",
                msecs=2000,
            )
            return True
        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export annotations to Word:\n{str(e)}",
            )
            return False

    def export_annotations_to_pdf(self, parent: QWidget | None = None) -> bool:
        """Export simple annotations to PDF via intermediate Word conversion."""
        try:
            if not self.session_data.is_file_open():
                QMessageBox.warning(
                    parent,
                    "No Session Data",
                    "No annotation file is currently open. Please open or create one first.",
                )
                return False

            if getattr(self.session_data, "tsv_file", None) is not None:
                try:
                    self.session_data.tsv_file.flush()
                except Exception:
                    pass

            from PySide6.QtWidgets import QFileDialog
            default_filename = self._generate_bids_report_filename('.pdf')
            start_dir = os.path.dirname(getattr(self.session_data, "file_path", "") or "")
            start_path = os.path.join(start_dir, default_filename) if start_dir else default_filename
            pdf_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Export Annotations Report",
                start_path,
                "PDF Files (*.pdf);;All Files (*)",
            )
            if not pdf_path:
                return False
            if not pdf_path.endswith(".pdf"):
                pdf_path += ".pdf"

            docx_path = os.path.splitext(pdf_path)[0] + "_tmp.docx"
            ok = self._export_annotations_to_word_path(docx_path)
            if not ok:
                QMessageBox.warning(
                    parent,
                    "No Data to Export",
                    "No annotations have been recorded yet.",
                )
                return False

            try:
                self._convert_docx_to_pdf(docx_path, pdf_path)
            finally:
                try:
                    os.unlink(docx_path)
                except Exception:
                    pass

            self._show_transient_message(
                parent,
                "Export Completed",
                f"Report saved successfully:\n{pdf_path}",
                msecs=2000,
            )
            return True
        except Exception as e:
            QMessageBox.critical(
                parent,
                "Export Error",
                f"Failed to export annotations to PDF:\n{str(e)}",
            )
            return False
